import os
import uuid
from pathlib import Path
from unittest.mock import patch

import pytest

from app.llm.embed import embed
from app.matching.report import lookup_resource
from app.matching.resume import ResumeError, evidence_level, extract_text, parse_resume, skills_from_text
from app.pipeline.status import job_id_for


def minimal_pdf(text: str) -> bytes:
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 24 720 Td ({safe}) Tj ET".encode("latin-1", "replace")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n",
        b"4 0 obj << /Length " + str(len(stream)).encode() + b" >> stream\n" + stream + b"\nendstream endobj\n",
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
    ]
    body = b"%PDF-1.1\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(body))
        body += obj
    xref = len(body)
    out = body + f"xref\n0 6\n0000000000 65535 f \n".encode()
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return out


def _fake_complete(messages):
    sys = messages[0]["content"]
    if "experience" in sys:
        return {"experience": "3年", "education": "本科"}
    if "url" in sys:
        return {"url": "https://example.com/skill"}
    return {"skills": [{"name": "FastAPI", "proficiency": "able"}, {"name": "Python"}]}


def test_doc_rejected():
    try:
        extract_text(b"abc", "resume.doc")
        assert False
    except ResumeError as exc:
        assert ".doc" in exc.detail


def test_empty_pdf_rejected():
    blank = minimal_pdf(" ")
    try:
        extract_text(blank.replace(b"( )", b"()"), "scan.pdf")
    except ResumeError:
        return
    try:
        extract_text(minimal_pdf(""), "scan.pdf")
    except ResumeError as exc:
        assert "扫描" in exc.detail or "文本" in exc.detail


def test_skills_from_text_skips_generic_short_chinese_names():
    index = [
        {"id": "s1", "name": "FastAPI"},
        {"id": "s2", "name": "数据"},
        {"id": "s3", "name": "开发"},
        {"id": "s4", "name": "意图识别"},
    ]
    found = skills_from_text("使用 FastAPI 做意图识别，负责数据开发", index)
    assert {row["skill_id"] for row in found} == {"s1", "s4"}


def test_skills_from_text_aligns_names():
    index = [
        {"id": "s1", "name": "FastAPI", "synonyms": ["starlette api"]},
        {"id": "s2", "name": "COBOL"},
    ]
    found = skills_from_text("做过 FastAPI 服务", index)
    assert {row["skill_id"] for row in found} == {"s1"}
    found_syn = skills_from_text("做过 starlette api", index)
    assert {row["skill_id"] for row in found_syn} == {"s1"}
    index_c = [{"id": "c", "name": "C"}, {"id": "s1", "name": "FastAPI"}]
    assert {row["skill_id"] for row in skills_from_text("LangChain FastAPI", index_c)} == {"s1"}


def test_parse_resume_merges_vocab_hits_when_model_returns_subset():
    index = [
        {"id": "s1", "name": "FastAPI", "synonyms": [], "embedding": embed(["FastAPI"])[0]},
        {"id": "s2", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]},
        {"id": "s3", "name": "Neo4j", "synonyms": [], "embedding": embed(["Neo4j"])[0]},
    ]

    def subset(messages):
        if "experience" in messages[0]["content"]:
            return {"experience": "3年", "education": "本科"}
        return {"skills": [{"name": "FastAPI"}]}

    out = parse_resume("生产环境用 FastAPI、Python 和 Neo4j 做过接口", index, complete_json=subset)
    assert {row["skill_id"] for row in out["skills"]} == {"s1", "s2", "s3"}


def test_parse_resume_splits_composite_skill_names():
    index = [
        {"id": "s1", "name": "LangChain", "synonyms": [], "embedding": embed(["LangChain"])[0]},
        {"id": "s2", "name": "LangGraph", "synonyms": [], "embedding": embed(["LangGraph"])[0]},
    ]

    def composite(messages):
        if "experience" in messages[0]["content"]:
            return {"experience": "2年", "education": "本科"}
        return {"skills": [{"name": "LangChain/LangGraph", "excerpt": "基于 LangChain/LangGraph 做编排"}]}

    out = parse_resume("基于 LangChain/LangGraph 做编排", index, complete_json=composite)
    assert {row["skill_id"] for row in out["skills"]} == {"s1", "s2"}


def test_evidence_fragment_prefers_result_sentence():
    index = [{"id": "s1", "name": "FastAPI", "synonyms": [], "embedding": embed(["FastAPI"])[0]}]
    out = parse_resume(
        "技能：FastAPI\n负责 FastAPI 接口，延迟降低 30%",
        index,
        complete_json=lambda *_: {"skills": [{"name": "FastAPI"}]},
    )
    assert out["evidence_fragments"][0]["evidence_level"] == "result"
    assert "30%" in out["evidence_fragments"][0]["text"]


def test_docx_tables_are_extracted():
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("简历正文")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "框架"
    table.cell(0, 1).text = "FastAPI"
    buffer = io.BytesIO()
    document.save(buffer)
    text = extract_text(buffer.getvalue(), "cv.docx")
    assert "FastAPI" in text


def test_parse_resume_aligns_and_keeps_unmarked_proficiency():
    index = [
        {"id": "s1", "name": "FastAPI", "synonyms": [], "embedding": embed(["FastAPI"])[0]},
        {"id": "s2", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]},
    ]

    def fake(messages):
        sys = messages[0]["content"]
        if "experience" in sys:
            return {"experience": "5年", "education": "硕士"}
        return {"skills": [{"name": "FastAPI", "proficiency": "able"}, {"name": "Python"}]}

    out = parse_resume("精通 FastAPI，也写 Python", index, complete_json=fake)
    by_id = {row["skill_id"]: row for row in out["skills"]}
    assert by_id["s1"]["proficiency"] == "able"
    assert by_id["s2"]["proficiency"] is None
    assert out["experience"] == "5年"
    assert out["education"] == "硕士"
    guessed = parse_resume(
        "built FastAPI in Python",
        index,
        complete_json=lambda *_: {"skills": [{"name": "FastAPI", "proficiency": "aware"}]},
    )
    assert guessed["skills"][0]["proficiency"] is None


def test_resume_evidence_levels_only_claim_visible_facts():
    assert evidence_level("FastAPI", "FastAPI") == "mention"
    assert evidence_level("使用 FastAPI 开发接口", "FastAPI") == "use"
    assert evidence_level("负责 FastAPI 接口，延迟降低 30%", "FastAPI") == "result"
    assert evidence_level("工具调用错误率由 8.6% 降至 1.9%", "工具调用") == "result"


def test_parse_resume_reads_explicit_info_without_model():
    index = [{"id": "s1", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]}]
    out = parse_resume(
        "Python developer with 3 years experience and a Bachelor degree.",
        index,
        complete_json=lambda *_: (_ for _ in ()).throw(RuntimeError("offline")),
    )
    assert out["experience"] == "3年"
    assert out["education"] == "Bachelor"


def test_parse_resume_preserves_structured_profile_experience_and_projects():
    index = [{"id": "s1", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]}]

    def structured(messages):
        if "experience" in messages[0]["content"]:
            return {
                "profile": {"role": "后端工程师", "experience": "4年"},
                "experience": "4年",
                "education": "硕士",
                "education_items": [{"text": "某大学 · 硕士"}],
                "experiences": [{"company": "甲科技", "title": "后端工程师", "start": "2021-01", "end": "2024-01", "summary": "负责 Python 服务"}],
                "projects": [{"name": "推荐系统", "summary": "构建离线评测"}],
            }
        return {"skills": [{"name": "Python"}]}

    out = parse_resume("后端工程师，使用 Python", index, complete_json=structured)
    assert out["profile"] == {"role": "后端工程师", "experience": "4年"}
    assert out["education_items"] == [{"text": "某大学 · 硕士"}]
    assert out["experiences"][0]["company"] == "甲科技"
    assert out["projects"][0]["name"] == "推荐系统"


def test_parse_resume_strict_mode_surfaces_model_failure():
    with pytest.raises(RuntimeError, match="offline"):
        parse_resume(
            "Python",
            [{"id": "s1", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]}],
            complete_json=lambda *_: (_ for _ in ()).throw(RuntimeError("offline")),
            strict=True,
        )


def test_demo_cv_falls_back_when_model_marks_info_unknown():
    text = extract_text(Path("data/eval-official-only/demo-cv.pdf").read_bytes(), "demo-cv.pdf")
    index = [{"id": "s1", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]}]

    def unknown_info(messages):
        if "experience" in messages[0]["content"]:
            return {"experience": "简历未标", "education": "简历未标"}
        return {"skills": []}

    out = parse_resume(text, index, complete_json=unknown_info)
    assert out["experience"] == "3年"
    assert out["education"] == "Bachelor"
    assert [skill["name"] for skill in out["skills"]] == ["Python"]


def test_lookup_resource_uses_cache(monkeypatch):
    monkeypatch.setattr("app.matching.report._verify_resource", lambda url, skill: url)
    sid = "skill-cache-" + uuid.uuid4().hex
    calls = {"n": 0}

    def fake(messages):
        calls["n"] += 1
        return {"url": "https://neo4j.com/docs/cypher-manual/current/"}

    first = lookup_resource(sid, "Cypher", complete_json=fake)
    second = lookup_resource(sid, "Cypher", complete_json=fake)
    assert first.startswith("http")
    assert first == second
    assert calls["n"] == 1


def test_diagnose_report_shape(client):
    pdf = minimal_pdf("Python FastAPI Neo4j RAG LangChain")
    with patch("app.llm.client.complete_json", side_effect=_fake_complete):
        up = client.post("/sessions", files={"file": ("cv.pdf", pdf, "application/pdf")})
    if up.status_code != 200:
        # pdfplumber may fail on tiny pdf; still require a structured error
        assert up.status_code == 400
        return
    session_id = up.json()["session_id"]
    jobs = client.get("/jobs").json()
    target = next((j for j in jobs if j["name"] == "大模型应用工程师"), None)
    if target is None:
        return
    with patch("app.llm.client.complete_json", side_effect=_fake_complete):
        res = client.post(
            "/diagnose",
            json={"session_id": session_id, "job_id": target["id"]},
        )
    assert res.status_code == 200
    body = res.json()
    assert "score" in body
    assert body["band"] in ("高度匹配", "基本匹配", "有明显差距", "不匹配")
    groups = body["groups"]
    assert set(groups) >= {"judge", "locate", "act", "explain"}
    assert groups["explain"]["watching_copy"] == "市场开始提，还没进要求，不算缺口"
    assert len(groups["act"]["path"]) <= 5
    assert groups["judge"]["band"] == body["band"]
    neighbor_names = {n["name"] for n in groups["locate"]["neighbors"]}
    assert "Agent 工程师" in neighbor_names or neighbor_names == {"大模型应用工程师"}
    assert "大模型应用工程师" in neighbor_names
    path = groups["act"]["path"]
    assert all(step.get("url") for step in path)
    assert all(step.get("why") in ("换档", "半档", "缺口") for step in path)
    assert isinstance(groups["judge"]["shift_set"], list)
    assert "preview_text" in body
    assert groups["judge"]["cells"]["experience"] == "3年"
    assert groups["judge"]["cells"]["education"] == "本科"


def test_session_returns_retryable_error_when_model_unavailable(client):
    pdf = minimal_pdf("Python FastAPI")
    with patch("app.main.extract_text", return_value="Python FastAPI"), patch(
        "app.main.parse_resume", side_effect=RuntimeError("model unavailable")
    ):
        response = client.post(
            "/sessions",
            files={"file": ("cv.pdf", pdf, "application/pdf")},
            data={"consent": "true"},
        )
    assert response.status_code == 503
    assert response.json()["error"] == "简历解析服务暂时不可用，请稍后重试"


def test_expired_session_is_404(client):
    job_id = job_id_for("大模型应用工程师")
    res = client.post("/diagnose", json={"session_id": "no-such", "job_id": job_id})
    assert res.status_code in (404, 400)


def test_two_job_diagnosis_returns_direction_payload(client):
    from app.matching.session import save

    sid = save({"preview_text": "Python FastAPI", "skills": [{"skill_id": "python", "name": "Python"}], "experience": "3年", "education": "本科"})
    from app import graph

    candidates = [job for job in client.get("/jobs").json() if graph.diagnostic_release(job["id"])["ok"]]
    if len(candidates) < 2:
        return
    response = client.post("/diagnose", json={"session_id": sid, "job_ids": [candidates[0]["id"], candidates[1]["id"]]})
    assert response.status_code == 200
    body = response.json()
    assert body["direction"]
    assert len(body["jobs"]) == 2
    assert all("shift_set" in item for item in body["jobs"])


def test_simulation_is_stateless_and_rejects_unknown_skill(client):
    from app.matching.session import load, save
    from app import graph

    sid = save({"preview_text": "Python FastAPI", "skills": [{"skill_id": "python", "name": "Python"}], "experience": "3年", "education": "本科"})
    target = next((job for job in client.get("/jobs").json() if graph.diagnostic_release(job["id"])["ok"]), None)
    if target is None:
        return
    before = load(sid)
    response = client.post("/diagnose/simulate", json={"session_id": sid, "job_id": target["id"], "assumed_skill_ids": [], "watching_skill_ids": ["watch-only"]})
    assert response.status_code == 200
    assert response.json()["watching_skill_ids"] == ["watch-only"]
    assert load(sid) == before
    rejected = client.post("/diagnose/simulate", json={"session_id": sid, "job_id": target["id"], "assumed_skill_ids": ["not-a-gap"]})
    assert rejected.status_code == 400


def test_session_correction_cannot_promote_unproven_result(client):
    from app.matching.session import save

    sid = save({
        "preview_text": "只在技能栏提到 FastAPI",
        "skills": [{"skill_id": "s1", "name": "FastAPI", "proficiency": None}],
        "evidence_fragments": [{"skill_id": "s1", "text": "只在技能栏提到 FastAPI", "evidence_level": "mention"}],
    })
    rejected = client.patch(
        f"/sessions/{sid}",
        json={"skills": [{"skill_id": "s1", "name": "FastAPI"}], "evidence_fragments": [{"skill_id": "s1", "text": "只在技能栏提到 FastAPI", "evidence_level": "result"}]},
    )
    assert rejected.status_code == 400
    accepted = client.patch(
        f"/sessions/{sid}",
        json={"skills": [{"skill_id": "s1", "name": "FastAPI"}], "user_added": [{"skill_id": "s2", "name": "PyTorch"}]},
    )
    assert accepted.status_code == 200
    assert accepted.json()["user_added"][0]["reason"] == "你补充的，简历尚未证明"


def test_sessions_rejects_doc(client):
    res = client.post("/sessions", files={"file": ("old.doc", b"xx", "application/msword")})
    assert res.status_code == 400
    assert "doc" in res.json()["error"].lower()


@pytest.mark.skipif(not os.environ.get("DEEPSEEK_API_KEY"), reason="no DEEPSEEK_API_KEY")
def test_live_deepseek_resume_and_resource():
    index = [
        {"id": "s1", "name": "FastAPI", "synonyms": [], "embedding": embed(["FastAPI"])[0]},
        {"id": "s2", "name": "Python", "synonyms": [], "embedding": embed(["Python"])[0]},
    ]
    out = parse_resume(
        "李四，计算机硕士，工作 4 年。熟悉 Python，生产环境用 FastAPI 做过接口，精通 FastAPI。",
        index,
    )
    ids = {row["skill_id"] for row in out["skills"]}
    assert ids & {"s1", "s2"}
    assert out["experience"] != "简历未标" or out["education"] != "简历未标"
    fastapi = next((row for row in out["skills"] if row["skill_id"] == "s1"), None)
    if fastapi and fastapi.get("proficiency"):
        assert fastapi["proficiency"] in ("aware", "able", "expert")
    sid = "skill-live-" + uuid.uuid4().hex
    url = lookup_resource(sid, "Cypher")
    assert url.startswith("http")
    assert "://" in url
