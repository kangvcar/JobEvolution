from __future__ import annotations

import io
import hashlib
import re
import subprocess
import sys
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from app.pipeline.align import align_skill, split_composite
from app.pipeline.extract import _PROF, _alias


class ResumeError(ValueError):
    def __init__(self, detail: str):
        super().__init__(detail)
        self.detail = detail


def extract_text(data: bytes, filename: str) -> str:
    if len(data) > 10 * 1024 * 1024:
        raise ResumeError("文件不能超过 10 MB")
    try:
        result = subprocess.run([sys.executable, "-m", "app.matching.resume", filename],
                                input=data, capture_output=True, timeout=15, cwd=Path(__file__).resolve().parents[2])
    except subprocess.TimeoutExpired as exc:
        raise ResumeError("文档解析超时，请简化文件后重试") from exc
    if result.returncode:
        raise ResumeError(result.stderr.decode("utf-8", "replace")[-200:] or "文档无法读取或超过解析资源限制")
    return result.stdout.decode("utf-8")


def _extract_document(data: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".doc") and not name.endswith(".docx"):
        raise ResumeError("不支持 .doc，请上传 PDF 或 docx")
    if name.endswith(".docx"):
        return _docx_text(data)
    if name.endswith(".pdf") or data[:4] == b"%PDF":
        return _pdf_text(data)
    raise ResumeError("请上传 PDF 或 docx")


def _pdf_text(data: bytes) -> str:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        if len(pdf.pages) > 30:
            raise ResumeError("PDF 不能超过 30 页")
        parts = [(page.extract_text() or "") for page in pdf.pages]
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeError("扫描件没有文本层")
    if len(text) > 50_000:
        raise ResumeError("提取文本不能超过 50,000 字符")
    return text


def _docx_text(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    text = "\n".join(part for part in parts if part and part.strip()).strip()
    if not text:
        raise ResumeError("文档没有可提取的文本")
    if len(text) > 50_000:
        raise ResumeError("提取文本不能超过 50,000 字符")
    return text


def _name_in_text(name: str, blob: str) -> bool:
    needle = (name or "").casefold()
    haystack = (blob or "").casefold()
    if not needle:
        return False
    if re.search(r"[\u4e00-\u9fff]", needle):
        return needle in haystack
    return re.search(rf"(?<![a-z0-9_]){re.escape(needle)}(?![a-z0-9_])", haystack) is not None


_SENTENCE_SPLIT = re.compile(r"[\n。！!]+|(?<!\d)\.(?!\d)")


def _split_sentences(text: str) -> list[str]:
    return [part.strip() for part in _SENTENCE_SPLIT.split(text or "") if part.strip()]


def evidence_level(text: str, name: str) -> str:
    """Classify only what the sentence visibly proves."""
    sentence = next((part for part in _split_sentences(text) if _name_in_text(name, part)), "") or (text or "").strip()
    if not sentence:
        return "mention"
    has_result = bool(re.search(r"\d+(?:\.\d+)?\s*(?:%|ms|秒|万|qps|次)?", sentence, re.I)) and bool(
        re.search(r"负责|主导|交付|实现|提升|降低|下降|降至|压缩|built|led|delivered|improv|reduc", sentence, re.I)
    )
    if has_result:
        return "result"
    if re.search(r"负责|使用|开发|构建|维护|实现|参与|used|built|develop|worked", sentence, re.I):
        return "use"
    return "mention"


_LEVEL_RANK = {"mention": 0, "use": 1, "result": 2}


def _sentences_for(text: str, names: list[str]) -> list[str]:
    found = []
    for sentence in _split_sentences(text):
        if any(_name_in_text(name, sentence) for name in names if name):
            found.append(sentence)
    return found


def _evidence_fragments(text: str, skills: list[dict], index: list[dict]) -> list[dict]:
    by_id = {row.get("id"): row for row in index}
    fragments = []
    for skill in skills:
        vocab = by_id.get(skill.get("skill_id"), {})
        names = [vocab.get("name") or skill.get("name") or "", *(vocab.get("synonyms") or [])]
        candidates = _sentences_for(text, names)
        llm_excerpt = str(skill.get("llm_excerpt") or "").strip()
        if llm_excerpt and llm_excerpt in text:
            candidates.append(llm_excerpt)
        if not candidates:
            continue
        candidates.sort(key=lambda sentence: (_LEVEL_RANK.get(evidence_level(sentence, names[0]), 0), len(sentence)), reverse=True)
        excerpt = candidates[0]
        fragments.append(
            {
                "id": "resume-evidence-" + hashlib.sha256(f"{skill['skill_id']}:{excerpt}".encode()).hexdigest()[:16],
                "skill_id": skill["skill_id"],
                "text": excerpt,
                "section": "project" if re.search(r"项目|project", excerpt, re.I) else "experience",
                "evidence_level": evidence_level(excerpt, names[0]),
            }
        )
    return fragments


INFO_PROMPT = (
    "Extract resume JSON. Fields: experience, education, profile, education_items, experiences, projects. "
    "experience is a short years string like 3年, or 简历未标. "
    "education is a short school/degree string, or 简历未标. "
    "profile may contain role and experience. education_items is a list of {text}. "
    "experiences is a list of {company,title,start,end,summary}; projects is a list of {name,summary}. "
    "Keep only facts visible in the resume, omit unknown fields, and do not invent."
)
SKILL_PROMPT = (
    "Extract resume skills JSON: {skills:[{name, proficiency, excerpt}]}. "
    "List every technical skill, tool, framework, platform, protocol, and domain method visible in the resume. "
    "Split composite names such as LangChain/LangGraph into separate items. "
    "excerpt is the shortest resume sentence that shows the skill; copy it verbatim. "
    "Omit proficiency unless the resume explicitly marks a level "
    "(了解/aware, 熟练/able, 精通/expert). Do not guess. "
    "Prefer names from this vocabulary when the resume uses a synonym.\n"
    "vocabulary: "
)


def parse_resume(
    text: str,
    index: list[dict],
    complete_json=None,
    *,
    threshold: float | None = None,
    strict: bool = False,
) -> dict:
    if complete_json is None:
        from app.llm.client import complete_json as complete_json
    blob = re.sub(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}|(?<!\d)1[3-9]\d{9}(?!\d)", "[联系方式已隐藏]", text or "")
    vocab = "、".join(row.get("name") or "" for row in index if row.get("name"))
    info_messages = [
        {"role": "system", "content": INFO_PROMPT},
        {"role": "user", "content": blob},
    ]
    skill_messages = [
        {"role": "system", "content": SKILL_PROMPT + vocab},
        {"role": "user", "content": blob},
    ]

    def _call(messages):
        try:
            payload = complete_json(messages)
            return payload if isinstance(payload, dict) else {}
        except Exception:
            if strict:
                raise
            return {}

    with ThreadPoolExecutor(max_workers=2) as pool:
        info_f = pool.submit(_call, info_messages)
        skill_f = pool.submit(_call, skill_messages)
        info = info_f.result()
        raw_skills = skill_f.result()
    skills = _align_skills(raw_skills.get("skills") or [], index, threshold=threshold)
    # 模型结果可以是子集，正文中的明确词表命中需要合并进来，再统一做证据校验。
    skills = _merge_skills(skills, skills_from_text(text, index, threshold=threshold))
    # 词表命中不能直接证明技能。只保留能定位到简历原文的项。
    grounded = {row["skill_id"] for row in _evidence_fragments(text, skills, index)}
    skills = [row for row in skills if row["skill_id"] in grounded]
    for row in skills:
        if not _marks_level_for_skill(text, row.get("name") or ""):
            row["proficiency"] = None
    fallback = _resume_info_from_text(text)
    profile_raw = info.get("profile") if isinstance(info.get("profile"), dict) else {}
    experience = str(info.get("experience") or "").strip()
    education = str(info.get("education") or "").strip()
    unknown = {"简历未标", "unknown", "未标注", "未提供", "n/a", "na"}
    experience = ("" if experience.casefold() in unknown else experience) or fallback["experience"] or "简历未标"
    education = ("" if education.casefold() in unknown else education) or fallback["education"] or "简历未标"
    role = str(profile_raw.get("role") or info.get("role") or "").strip()
    profile_experience = str(profile_raw.get("experience") or "").strip()
    if profile_experience and profile_experience != "简历未标":
        experience = profile_experience
    education_items = _structured_items(info.get("education_items"), ("text",))
    if not education_items and education != "简历未标":
        education_items = [{"text": education}]
    experiences = _structured_items(info.get("experiences"), ("company", "title", "start", "end", "summary"))
    projects = _structured_items(info.get("projects"), ("name", "summary"))
    fragments = _evidence_fragments(text, skills, index)
    return {
        "profile": {"role": role, "experience": experience},
        "education": education,
        "education_items": education_items,
        "experiences": experiences,
        "projects": projects,
        "experience": experience,
        "education_text": education,
        "skills": skills,
        "evidence_fragments": fragments,
        "date_conflicts": date_conflicts(text),
        "user_added": [],
    }


def _structured_items(value, fields: tuple[str, ...]) -> list[dict]:
    if not isinstance(value, list):
        return []
    items = []
    for raw in value:
        if not isinstance(raw, dict):
            continue
        item = {field: str(raw[field]).strip() for field in fields if raw.get(field) not in (None, "")}
        if item:
            items.append(item)
    return items


def date_conflicts(text: str) -> list[dict]:
    # ponytail: bounded date regex and pairwise overlap scan; use a structured timeline parser if formats or volume expand.
    ranges = []
    for match in re.finditer(r"(20\d{2})[./-](\d{1,2}).{0,3}(20\d{2})[./-](\d{1,2})", text or ""):
        start = int(match.group(1)) * 12 + int(match.group(2))
        end = int(match.group(3)) * 12 + int(match.group(4))
        if end < start:
            ranges.append({"text": match.group(0), "reason": "结束时间早于开始时间"})
        ranges.append({"start": start, "end": end, "text": match.group(0)})
    conflicts = [row for row in ranges if "reason" in row]
    for i, left in enumerate(ranges):
        for right in ranges[i + 1 :]:
            if "start" in left and "start" in right and max(left["start"], right["start"]) <= min(left["end"], right["end"]):
                conflicts.append({"left": left["text"], "right": right["text"], "reason": "经历时间重叠"})
    return conflicts


def _marks_level_for_skill(text: str, name: str) -> bool:
    blob = text or ""
    needle = re.escape(name)
    return re.search(rf"(?:了解|熟练|精通|aware|able|expert|proficient)\s*(?:掌握|使用|过)?\W{{0,12}}{needle}", blob, re.I) is not None


def _resume_info_from_text(text: str) -> dict:
    # ponytail: explicit common formats only; use structured parsing for multilingual employment histories.
    blob = text or ""
    experience = re.search(r"\b(\d+)\s*(?:年|years?)\s*(?:工作|经验|experience)?", blob, re.I)
    education = re.search(r"\b(Bachelor|Master|PhD)\b|(?:博士|硕士|本科|大专)(?:学历|学位)?", blob, re.I)
    return {
        "experience": f"{experience.group(1)}年" if experience else "",
        "education": education.group(0) if education else "",
    }


def _skill_name_parts(name: str, index: list[dict]) -> list[str]:
    parts = []
    for chunk in re.split(r"[,，、;；|]", name or ""):
        chunk = chunk.strip()
        if chunk:
            parts.extend(split_composite(chunk, index))
    return parts or ([name] if name else [])


def _merge_skills(primary: list[dict], extra: list[dict]) -> list[dict]:
    found = []
    seen: set[str] = set()
    for row in (*primary, *extra):
        sid = str(row.get("skill_id") or "")
        if not sid or sid in seen:
            continue
        seen.add(sid)
        found.append(row)
    return found


def _align_skills(rows: list, index: list[dict], *, threshold: float | None = None) -> list[dict]:
    found = []
    seen: set[str] = set()
    for raw in rows:
        if isinstance(raw, str):
            raw = {"name": raw}
        if not isinstance(raw, dict):
            continue
        name = str(raw.get("name") or "").strip()
        if not name:
            continue
        marked = str(raw.get("proficiency") or "").strip()
        proficiency = _alias(_PROF, marked, "") if marked else ""
        if proficiency not in ("aware", "able", "expert"):
            proficiency = None
        excerpt = str(raw.get("evidence") or raw.get("excerpt") or "").strip()
        for part in _skill_name_parts(name, index):
            hit = align_skill(part, index, threshold=threshold)
            if hit is None or hit["id"] in seen:
                continue
            seen.add(hit["id"])
            found.append(
                {
                    "skill_id": hit["id"],
                    "name": hit.get("name") or part,
                    "proficiency": proficiency,
                    "llm_excerpt": excerpt,
                }
            )
    return found


def _usable_surface(name: str) -> bool:
    text = (name or "").strip()
    if not text:
        return False
    if re.search(r"[A-Za-z]{2,}", text) or "+" in text:
        return True
    return len(text) >= 3


def skills_from_text(text: str, index: list[dict], *, threshold: float | None = None) -> list[dict]:
    # ponytail: name+synonym substring; cosine via align_skill when resume phrases diverge from Skill.name
    blob = (text or "").casefold()
    found = []
    seen: set[str] = set()
    for skill in index:
        names = [n for n in [skill.get("name") or "", *(skill.get("synonyms") or [])] if _usable_surface(n)]
        if not any(_name_in_text(n, blob) for n in names):
            continue
        sid = skill["id"]
        if sid in seen:
            continue
        seen.add(sid)
        found.append(
            {
                "skill_id": sid,
                "name": skill.get("name") or names[0],
                "proficiency": None,
            }
        )
    return found


if __name__ == "__main__":
    import resource
    resource.setrlimit(resource.RLIMIT_CPU, (10, 10))
    if sys.platform == "linux":
        resource.setrlimit(resource.RLIMIT_AS, (768 * 1024 * 1024, 768 * 1024 * 1024))
    try:
        sys.stdout.write(_extract_document(sys.stdin.buffer.read(10 * 1024 * 1024 + 1), sys.argv[1]))
    except Exception as exc:
        sys.stderr.write(exc.detail if isinstance(exc, ResumeError) else "文档已损坏或格式无效，请重新导出 PDF 或 docx")
        sys.exit(1)
