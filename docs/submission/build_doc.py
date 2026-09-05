# -*- coding: utf-8 -*-
"""Build 作品设计实现方案 (Word) for XH-202621 智演 JobEvolution."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/kangvcar/Documents/code/JobEvolution/XH-202621_作品设计实现方案_智演JobEvolution.docx"
FIG = "/tmp/je_doc"
SHOT = "/tmp/je_shots"

BODY_FONT_CN = "宋体"
HEAD_FONT_CN = "黑体"
BODY_FONT_EN = "Times New Roman"

doc = Document()

# ------------------------------------------------------------------ page setup
for s in doc.sections:
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.6)

# update fields on open (TOC)
settings = doc.settings.element
upd = OxmlElement("w:updateFields")
upd.set(qn("w:val"), "true")
settings.append(upd)


def set_run_font(run, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=None, bold=None, color=None):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def style_font(style, cn, en, size, bold=False, color=(0, 0, 0)):
    style.font.name = en
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)


normal = doc.styles["Normal"]
style_font(normal, BODY_FONT_CN, BODY_FONT_EN, 12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(4)

for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12.5)):
    st = doc.styles[name]
    style_font(st, HEAD_FONT_CN, BODY_FONT_EN, size, bold=True)
    st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.3

for name in ("List Bullet", "List Number"):
    st = doc.styles[name]
    style_font(st, BODY_FONT_CN, BODY_FONT_EN, 12)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.space_after = Pt(2)

FIG_NO = [0]
TAB_NO = [0]


def H(level, text):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, cn=HEAD_FONT_CN, bold=True, color=(0, 0, 0))
    return p


def P(text, bold=False, italic=False, align=None, size=None, indent=True, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    r.italic = italic
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def PR(parts, indent=True):
    """paragraph with mixed runs: list of (text, bold)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    for text, bold in parts:
        r = p.add_run(text)
        set_run_font(r, bold=bold)
    return p


def BUL(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            head, rest = it
            r = p.add_run(head)
            set_run_font(r, bold=True)
            r2 = p.add_run(rest)
            set_run_font(r2)
        else:
            r = p.add_run(it)
            set_run_font(r)


def NUM(items):
    BUL(items, style="List Number")


def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    r.font.size = Pt(9.5)
    # shading
    ppr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F3F3")
    ppr.append(shd)
    return p


def shade(cell, fill):
    tcpr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def TABLE(caption, header, rows, widths=None, font_size=10.5):
    TAB_NO[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"表 {TAB_NO[0]}  {caption}")
    set_run_font(r, cn=HEAD_FONT_CN, size=10.5, bold=True)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        rr = p.add_run(h)
        set_run_font(rr, cn=HEAD_FONT_CN, size=font_size, bold=True)
        shade(c, "E7E6E6")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(str(v))
            set_run_font(rr, size=font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def FIGURE(path, caption, width_in=6.0):
    if not os.path.exists(path):
        P(f"[图缺失：{os.path.basename(path)}]", color=(200, 0, 0))
        return
    FIG_NO[0] += 1
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(f"图 {FIG_NO[0]}  {caption}")
    set_run_font(r, cn=HEAD_FONT_CN, size=10.5, bold=True)


def add_field(paragraph, instr):
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "（右键选择“更新域”以生成目录）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._element.append(fld_begin)
    r._element.append(instr_el)
    r._element.append(fld_sep)
    r._element.append(txt)
    r._element.append(fld_end)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# footer page number
for s in doc.sections:
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._element.append(b); r._element.append(it); r._element.append(e)
    set_run_font(r, size=9)

# ================================================================== cover
for _ in range(5):
    doc.add_paragraph()
P("2026 年“挑战杯”揭榜挂帅擂台赛", align="center", size=16, indent=False)
P("题目编号：XH-202621", align="center", size=14, indent=False)
P("发榜单位：科大讯飞股份有限公司", align="center", size=14, indent=False)
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("多源异构数据驱动岗位和能力图谱构建\n与动态演化分析研究"); set_run_font(r, cn=HEAD_FONT_CN, size=22, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("作品设计实现方案"); set_run_font(r, cn=HEAD_FONT_CN, size=26, bold=True)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("作品名称：智演 JobEvolution —— 多源招聘数据驱动的岗位能力图谱与职业迁移诊断系统"); set_run_font(r, size=13)
for _ in range(5):
    doc.add_paragraph()
for line in ("参赛单位：______________________", "团队成员：______________________", "指导教师：______________________", "提交日期：2026 年 9 月"):
    P(line, align="center", size=13, indent=False)
page_break()

# ================================================================== TOC
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
_r = tt.add_run("目  录"); set_run_font(_r, cn=HEAD_FONT_CN, size=16, bold=True)
tp = doc.add_paragraph()
add_field(tp, 'TOC \\o "1-3" \\h \\z \\u')
page_break()

# ================================================================== 1 概述
H(1, "1  项目概述")
H(2, "1.1  赛题背景与问题")
P("在国家从“人口红利”向“人才红利”转型、“人工智能+”行动深入实施的背景下，新一代信息技术领域的技术迭代速度已明显快于人才培养周期。劳动力市场呈现出典型的结构性矛盾：企业在新兴岗位上“招不到合适的人”，新兴领域青年则面临职业发展路径模糊、技能需求变化快的问题。传统的人才盘点与招聘手段依赖关键词匹配，缺乏对技术趋势的实时感知，无法回答“技术 A 的爆发会给新兴岗位 B 带来哪些新技能要求”这类动态关联问题。")
P("赛题要求围绕“数据驱动 + 大模型 + 知识图谱”方向，攻克三大核心难题：一是从多源数据中精准发现并定义新兴岗位、动态更新既有岗位的能力要求；二是解决招聘数据的“时滞”“噪声”与“抄袭”问题，并防控 AI 生成内容的“幻觉”；三是实现高精度的简历解析与细粒度的人岗差距分析，最终形成从“新岗位发现与定义”“既有岗位能力动态更新”到“全景图谱可视化”“人岗匹配诊断”和“改进建议与学习路径规划”的完整系统。")

H(2, "1.2  作品定位与核心价值")
P("本作品“智演 JobEvolution”是一套面向新一代信息技术领域（人工智能、大数据、智能系统、物联网四个固定领域）的岗位能力图谱构建与动态演化分析系统。系统持续采集企业官方招聘门户发布的岗位说明书（JD），以大模型完成结构化抽取，以知识图谱表达岗位、技能点、证据与演化事件之间的关系，以确定性统计规则和多重审核闸门约束模型输出，最终为求职者提供以证据为核心的人岗匹配诊断、差距分析与学习路径规划，为企业管理员提供可审计的岗位能力更新工作台。")
P("与传统静态岗位画像相比，本作品的核心差异体现在四个方面：")
BUL([
    ("证据优先。", "图谱中的每一条岗位要求边都能回溯到具体的 JD 原文摘录与来源公司，每一条演化事件都携带证据引用与审核状态，任何没有证据链的事实都不能进入正式图谱。"),
    ("双时间演化。", "系统分别记录“观察时间”（何时在数据源看到）与“有效期”（能力要求何时生效与失效），旧要求边写失效时间而非删除，从而以切片差分的方式表达岗位能力的动态演化，并天然解决 JD 数据的时滞问题。"),
    ("确定性先于模型。", "新岗位状态流转、要求性质判定、置信层、匹配档位、缺口集与换档条件全部由可复现的规则与统计计算得出，大模型只负责抽取、簇判别与结论解释，且解释必须引用已存在的证据。"),
    ("多层幻觉防控。", "从切段、强 schema 抽取、原文回指校验、三层技能归一、覆盖率入池、要求判定票、置信层、人工与独立模型审核，到诊断发布校验和不可变图谱发布版本，形成十道闸门，任何一道未通过只会降级为“观测中”“审核提案”或“暂停诊断”。"),
])

H(2, "1.3  与赛题要求的对应关系")
TABLE("赛题要求与系统功能对应表", ["赛题要求", "系统功能与实现", "所在模块 / 页面"], [
    ["岗位选择范围：新一代信息技术（AI、大数据、智能系统、物联网）", "图谱固定四个领域节点，采集阶段按标题领域词粗滤，数据深度优先人工智能", "collectors/domain.py；/graph 领域筛选"],
    ["①新岗位发现与定义", "嵌入聚类 + LLM 簇三分类（新岗位 / 别名 / 噪声），候选→萌芽→成型状态机由独立源计票自动流转；岗位定义以“定义声明”为单位逐条挂证据，支持人工批 / 改 / 驳与每日重算", "pipeline/discover.py、status.py、gate.py；/discover、/admin"],
    ["②既有岗位能力动态更新", "要求边带 valid_from / valid_to，每日管线产出新增、修改、失效三类演化事件并附数据源；管理页按岗位版本分组审核、批量批准、撤回证据", "pipeline/gate.py、graph.py period_delta；/graph 切片差分、/admin"],
    ["③新一代信息技术岗位全景图谱", "领域 / 岗位 / 技能类目 / 技能点四层，颗粒度到技能点；可按技术栈（五个技能类目）与适用级别（初 / 中 / 高）切换视图；拓扑图与表格双视图，点技能点打开证据抽屉", "web/app/graph（React Flow + d3-force Web Worker）"],
    ["④人岗匹配度诊断与差距分析", "PDF / docx 简历解析 → 技能点对齐 → 用户校对 → 岗位推荐序 → 双岗对照报告：匹配档位、缺口集、简历证据级、换档条件、双轨行动清单、学习路径", "matching/resume.py、score.py、bands.py、report.py；/diagnose 五步流程"],
    ["创新①多源异构数据清洗与交叉验证", "fingerprint 幂等 + 64 位 simhash 近重去重解决“抄袭”；独立源按规范化公司名计票；覆盖率随周期重算；要求判定票 60% 且≥2 独立源；要求等价数上限与岗位要求异常拦截“通胀”", "collectors/simhash.py、normalize.py；pipeline/gate.py、diagnostic_release.py"],
    ["创新②能力“幻觉”防控", "十道闸门（见第 9 章）；抽取模型不得审核自身输出；低置信永不自动入谱；报告中每条判断必须携带证据 ID", "pipeline/*、matching/report.py、docs/adr"],
    ["可验证性：≥100 条 JD 与测试用例，三项准确率可量化", "100 条 JD 金标、100 份简历金标、100 对匹配金标，阈值冻结于 freeze.json；set-based F1 自动评测；pytest 覆盖率 69.67%（≥60%）", "data/eval-official-only/、app/eval/、tests/、.github/workflows/ci.yml"],
    ["提交形式：源码、Docker 部署、单测、测试数据", "开源仓库 + docker compose 一键部署 + 22 个测试文件 180 个用例 + 两岗提交物（Agent 工程师 / 大模型应用工程师）", "docker-compose.yml、tests/、data/eval-official-only/deliver/"],
], widths=[4.2, 7.3, 4.0], font_size=9.5)

# ================================================================== 2 需求分析
H(1, "2  需求分析")
H(2, "2.1  用户与使用场景")
BUL([
    ("求职者（主路径）。", "已有 1 至 5 年后端 / 全栈经验、准备转向大模型应用或 Agent 岗位的工程师，以及在校学生和其他技术求职者。他们带一份含文本层的 PDF 或 docx 简历进入系统，免登录完成“上传 → 校对 → 选岗 → 生成 → 报告”五步，得到方向结论、缺口集、换档条件和学习路径。简历只保存在一小时有效的匿名会话中。"),
    ("企业 HR / 图谱管理员（管理路径）。", "通过口令进入管理后台，查看待审队列中的新岗位首次发布、核心必备技能新增、低置信抽取与技能合并提案，进行批准、改写后批准或驳回；按岗位版本批量审核；管理采集门户；查看运行状态与审核统计。"),
    ("市场观察者。", "不上传简历，仅通过首页总览、图谱工作台与市场演化页了解哪些岗位正在形成、哪些技能刚跨过入池门槛、哪些要求已失效。"),
])
H(2, "2.2  功能需求")
TABLE("功能需求清单", ["编号", "功能", "需求描述", "优先级"], [
    ["F1", "多源采集", "每日从企业官方招聘门户采集四领域 JD，去重落盘为 JD 快照，写入事件流；支持断点续采、增量提前停止与门户启停管理", "高"],
    ["F2", "结构化抽取", "以大模型将 JD 抽取为岗位名、领域、目标岗位、原子技能点列表（性质、熟练级、置信、原文摘录、段落、类目、候选类型），经 Pydantic 校验", "高"],
    ["F3", "新岗位发现与定义", "对未对齐岗位的 JD 聚类、判别、计票并流转状态；生成含名称、核心职责、必备技能、加分技能、典型应用场景的岗位定义；支持人工优化与动态更新", "高"],
    ["F4", "既有岗位能力更新", "识别既有岗位要求边的新增、修改、失效，生成带证据的演化事件与更新说明；支持人工批 / 改 / 驳、批量批准与撤回", "高"],
    ["F5", "全景图谱可视化", "领域 / 岗位 / 技能类目 / 技能点四层展示；按技术栈、级别、必备 / 加分、近 90 天生效切换；证据抽屉；切片差分", "高"],
    ["F6", "简历解析", "解析 PDF / docx，抽取基本信息、教育、经历、项目与技能点，标注简历证据片段与证据级；允许用户按原文修正", "高"],
    ["F7", "匹配诊断与差距分析", "匹配分、档位、缺口集、熟练级差异、换档条件；岗位推荐序；双岗方向结论；无法区分判定", "高"],
    ["F8", "改进建议与学习路径", "双轨行动清单（简历证明轨 / 能力提升轨）、按换档条件排序的学习路径、可验证的学习资源、换档模拟器、简历证据地图、邻近岗位迁移地图、市场信号雷达、求职叙事稿", "高"],
    ["F9", "审核与发布", "待审队列、置信层、自动审核开关、岗位版本批量审核、诊断发布校验、不可变发布版本与回滚", "高"],
    ["F10", "评测与测试", "三项准确率评测脚本、金标起草与裁决工具、单元测试与覆盖率门禁、CI", "高"],
], widths=[1.2, 2.8, 9.5, 1.6], font_size=9.5)
H(2, "2.3  非功能需求")
BUL([
    ("准确性。", "JD 解析、简历提取、人岗匹配三项均以技能点集合的 set-based F1 度量，目标≥0.90；评测阈值冻结、金标独立、不喂解析输出。"),
    ("可验证性。", "任何非平凡逻辑都有可运行检查；单元测试覆盖率≥60%；三项评测可一键复跑；报告中的数字全部来自未 mock 的本地脚本。"),
    ("隐私与安全。", "简历原文件解析后立即删除，只把提取文本发送给配置的模型服务商；会话 TTL 一小时；管理口令常量时间比较，会话 Cookie Secure / HttpOnly / SameSite=Strict；Cypher 全部参数化；上传校验 MIME 与大小；日志不记录简历正文与口令。"),
    ("可部署性。", "单台服务器 Docker Compose 一键启动 web、api、pipeline、neo4j、redis 五个容器；健康检查齐全；从图谱快照可在空卷上重建。"),
    ("性能与体验。", "图谱布局在 Web Worker 中计算，主线程不阻塞；图谱页首屏 JS 约 174 KB；诊断报告同步返回；键盘可达、焦点可见、支持 320px 视口与 200% 缩放。"),
    ("可迁移性。", "四领域框架固定，岗位、技能点与状态全部由管线从数据长出，不手录；新增领域或岗位只需扩充对齐靶子与采集门户。"),
])

# ================================================================== 3 总体设计
H(1, "3  总体设计")
H(2, "3.1  设计原则")
NUM([
    "数据可信先于报告文案：前端不接一套仍会变化的要求与诊断契约，先把技能、要求边与岗位版本清理干净。",
    "确定性计算先于模型解释：岗位排序、要求性质、档位、换档条件与“无法区分”全部由结构化规则决定，模型只解释。",
    "所有职业判断能回到证据：引用已有事实时带证据 ID，缺失时只写“简历中未找到”，不推断求职者真实能力。",
    "共享函数只改一处：技能对齐 align_skill、置信层、状态机、匹配分各只有一个入口，匹配与换档模拟复用同一套服务端函数。",
    "单岗失败只隔离单岗：某岗位未通过诊断发布校验时只将其排除出推荐与诊断，不阻断同版本其他岗位。",
    "撤回不是演化：从未成立或证据不足的事实作废时不进入切片差分、演化趋势和匹配，只有真实变化才关闭旧事实有效期。",
])
H(2, "3.2  系统架构")
P("系统采用“数据源 → 采集 → 图谱构建 → 存储 → 服务 → 应用”六层架构，如图所示。求职者路径只读取图谱发布版本和自身会话，不接触采集与审核；管理员路径通过口令门进入待审队列与运维状态。所有大模型调用经由唯一出口 app/llm/client.py，按环境变量在 DeepSeek、B.AI、Tuzi 三个 OpenAI 兼容供应商之间切换；嵌入统一使用 BAAI/bge-m3，无嵌入密钥时回落本地哈希向量以保证测试与 CI 不出网。")
FIGURE(f"{FIG}/fig_arch.png", "系统总体架构", 6.2)
H(2, "3.3  技术选型")
TABLE("技术选型与理由", ["层次", "选型", "选型理由"], [
    ["后端", "Python 3.12 + FastAPI + Pydantic", "异步 HTTP、类型化模型校验，LLM JSON 输出可直接落到 Pydantic 校验；与采集、抽取、评测脚本共用一套代码"],
    ["前端", "Next.js 15（App Router）+ React 19 + TypeScript", "五个路由、服务端渲染与静态生成；类型安全；next/font 加载 IBM Plex 字体"],
    ["图谱可视化", "@xyflow/react（React Flow）+ d3-force + dagre", "节点 / 边渲染与交互成熟；力导向布局在 Web Worker 中计算避免主线程阻塞；防碰撞保证零重叠"],
    ["图数据库", "Neo4j 5 Community", "原生属性图，边上直接携带有效期、置信、来源列表；Cypher 表达切片查询简洁；单容器即可满足规模"],
    ["缓存 / 会话 / 事件", "Redis 7（AOF 持久化）", "简历会话 TTL、采集指纹集合、事件流（Stream）、学习资源缓存、直通开关，一套组件解决"],
    ["大模型", "DeepSeek V4 Flash / B.AI / Tuzi GPT-5.6（OpenAI 兼容）", "非思考模式 + JSON 模式，低延迟结构化抽取；供应商可切换，避免单点依赖；额度与费用有硬上限"],
    ["嵌入", "BAAI/bge-m3（硅基流动 OpenAI 兼容端点）", "中英混排技能名向量化效果好；用于技能对齐、实体消解与岗位聚类；阈值按金标校准"],
    ["简历解析", "pdfplumber + python-docx", "取文本层即可满足含文本层的 PDF / docx；不引入 OCR 与重型解析框架"],
    ["聚类", "scikit-learn 层次聚类（最小簇 3）", "小规模簇发现足够稳定，不为演示引入更重的 HDBSCAN"],
    ["部署", "Docker Compose", "单服务器五容器，一条命令启动；测试库与产品库数据卷隔离"],
    ["测试", "pytest + pytest-cov + GitHub Actions", "覆盖率门禁 60%；CI 以 mock 模型跑三项评测与全部单测"],
], widths=[2.6, 5.4, 7.5], font_size=9.5)

H(2, "3.4  图谱本体设计")
P("图谱包含六类节点与五类主要关系。岗位不按初级 / 高级分裂节点，级别差异表达在要求边的 levels 属性上；证据只作为节点存在，其 ID 写入要求边的 sources 列表与演化事件的 payload，保证 Community 版 Neo4j 也能表达“边引用证据”。")
FIGURE(f"{FIG}/fig_ontology.png", "岗位能力图谱本体", 6.0)
TABLE("节点定义", ["标签", "关键属性", "说明"], [
    ["Domain 领域", "id ∈ {ai, data, system, iot}，name", "新一代信息技术四个固定顶层分区，代码中唯一写死的节点"],
    ["Job 岗位", "id，name，status（candidate / emerging / formed），大典编码、esco_id、onet_id（可空），watching[]", "一类职位的规范化表示；watching 保存已抽出但未达入池门槛的“观测中”技能点"],
    ["SkillCategory 技能类目", "id，name ∈ {语言, 框架, 平台, 工程, 领域知识}", "全景图谱“技术栈视图”的聚合层，只用于导航，不参与差距分析对账"],
    ["Skill 技能点", "id，name，同义词列表，embedding_ref", "可由职责或要求原文证明的原子技术技能或领域知识，是差距分析的对账单位"],
    ["Evidence 证据", "id，path，source，company（规范化公司名），observed_at，simhash，retracted", "一条去重后的 JD 快照；独立源计票只认 company"],
    ["EvolutionEvent 演化事件", "id，kind，at，confidence，review（pending / approved / auto_passed / rejected），payload", "图谱变更的原子记录，分边级（要求新增 / 移除 / 修改）与节点级（新岗位入谱 / 定义更新 / 状态流转）；待审队列即 review=pending 的事件"],
    ["RequirementVersion / RequirementGroup / GraphRelease", "版本化要求、要求组（min_required）、图谱发布版本", "支撑版本化事实、要求组对账与不可变发布"],
], widths=[3.4, 6.0, 6.1], font_size=9.5)
TABLE("关系定义", ["关系", "方向", "关键属性 / 语义"], [
    ["IN_DOMAIN", "Job → Domain", "岗位所属领域"],
    ["IN_CATEGORY", "Skill → SkillCategory", "技能点所属技术栈类目"],
    ["REQUIRES", "Job → Skill", "kind（required / bonus）、proficiency（aware / able / expert）、weight、levels[]、valid_from、valid_to、confidence、layer（high / mid / low）、sources[]（证据 ID）、group_id / min_required"],
    ["AFFECTS", "EvolutionEvent → Job", "事件影响的岗位；payload 中记录 skill_id 与新旧边字段"],
    ["ALIAS_OF", "Job → Job", "簇判别为别名时并入既有岗位，如“大模型应用开发工程师”→“大模型应用工程师”"],
], widths=[3.0, 3.6, 8.9], font_size=9.5)
P("某一时点的岗位切片查询只需一条 Cypher：")
CODE("MATCH (j:Job {id:$job})-[r:REQUIRES]->(s:Skill)\nWHERE r.valid_from <= $t AND (r.valid_to IS NULL OR $t < r.valid_to)\nRETURN j, r, s")

H(2, "3.5  数据流总览")
P("每日管线由独立的 pipeline 容器串行执行“采集 → 抽取 → 消解对齐 → 入池计票 → 置信分层 → 待审或自动审核 → 合并演化事件 → 诊断发布校验 → 原子发布图谱版本”。求职者的诊断会话独立于管线：上传简历后创建一小时有效的 Redis 会话，诊断时固定读取创建时的图谱发布版本，避免报告中途因发布而漂移。")
CODE("官方招聘门户 ─fingerprint 幂等─▶ data/official-only/jd/{id}.json + Redis SET ingest:fp\n        │  XADD jobs:events（采集进度）\n        ▼\n抽取 worker ── LLM JSON ──▶ Pydantic ──▶ 原文回指校验\n        ├─ 技能对齐（表面归一 → 同义词 → bge 余弦）\n        ├─ 入池：职责 / 要求段 ∧ 簇内覆盖率 ≥30%\n        ├─ 要求判定票：明确票 ≥60% ∧ ≥2 独立源\n        └─ 置信层 → 待审队列 / 独立模型自动审核\n                │\n                ▼\n         Neo4j 主图 ── 诊断发布校验 ──▶ 图谱发布版本（不可变，可回滚）\n                │\nNext.js ◀── FastAPI ◀── 简历会话（Redis TTL 1h）")

# ================================================================== 4 采集
H(1, "4  多源异构数据采集与清洗")
H(2, "4.1  数据源")
P("系统以企业官方招聘门户作为主数据源，而非第三方招聘网站的二次转载。官方门户的 JD 由企业直接发布，发布时间可信、正文完整、不存在平台代理商的批量复制，是解决“抄袭”与“时滞”问题的第一道保障。当前配置 54 个门户（53 个启用），覆盖智谱、月之暗面、MiniMax、小米、蔚来、小鹏、理想、字节跳动、腾讯、生数科技、深言、众擎机器人、轻舟智航、鉴智机器人、叠纸游戏、蓝箭航天、后摩智能等企业，门户类型包括飞书招聘（51 个）、腾讯招聘、字节跳动招聘与北森。累计去重 JD 快照 3,595 条，覆盖人工智能、大数据、智能系统、物联网四个领域。")
TABLE("采集门户类型与适配方式", ["门户类型", "数量", "接口形态", "适配要点"], [
    ["飞书招聘（jobs.feishu.cn）", "51", "公开 JSON 列表 + 详情接口", "自动解析门户路径；分页扫描；按关键词命中四领域标题；过滤实习岗"],
    ["腾讯招聘", "1", "公开 JSON", "关键词搜索 + 分页；正文 HTML 去标签"],
    ["字节跳动招聘", "1", "公开 JSON", "分类筛选 + 分页"],
    ["北森（Beisen）", "1", "租户 JSON", "租户 host 参数化；字段映射到统一记录"],
], widths=[4.2, 1.4, 4.0, 5.9], font_size=9.5)
H(2, "4.2  采集架构与增量策略")
P("采集模块分为三层：source 负责按门户类型产出统一的原始记录（公司、岗位名、正文、发布时间、渠道、来源 URL），controller 负责去重、落盘与证据写入，sink 负责向 Redis Stream 写入采集事件。最多 8 个 worker 并发抓取门户，主线程按门户完成顺序串行写快照与证据，避免跨线程近重复竞争；每份快照原子写入磁盘，每个门户完成后刷新证据。")
BUL([
    ("断点续采。", "collect.checkpoint.json 记录每个门户的完成点，进程中断后再次执行会跳过已完成门户，未完成门户安全重放；collect.lock 防止并发采集。"),
    ("增量提前停止。", "每日扫描每个关键词从 offset=0 开始，当连续两个分页的岗位都已由 Redis 正文指纹确认未变化时提前停止；需要补历史数据时设置 COLLECT_FULL_SCAN=1。"),
    ("抽取缓存与检查点。", "每份成功抽取的 JD 立即写入 .extract-v4.cache 与 extract_completed 检查点，失败样本下次只重试未完成项；抽取最多 3 次尝试，仍失败则进入待审并标记抽取失败。"),
    ("门户管理。", "管理后台可新增、启停、删除门户，触发一次采集，并通过 SSE 实时查看 collect_started、jd_ingested、collect_portal_failed、collect_finished 等事件。"),
])
H(2, "4.3  去重、幂等与“抄袭”识别")
BUL([
    ("站点级幂等。", "fingerprint = sha256(source + job_id + 正文 simhash)；没有站点 ID 时使用规范化的“公司 | 标题 | 城市”。命中 Redis 集合 ingest:fp 的记录直接跳过。"),
    ("正文级近重去重。", "对正文计算 64 位 simhash，Hamming 距离≤3 视为近重复。近重复的 JD 只保留观察时间最早的一条作为证据，其余不计入独立源。同一公司在多个门户重复发布、不同公司互相抄袭 JD 模板，都会被折叠为一票。"),
    ("公司规范化。", "去除“有限公司 / 股份 / 括号地名”等后缀得到规范化公司名，独立源按规范化公司名去重计票；渠道名（如飞书、北森）不计为独立源。"),
    ("正文完整性。", "缺正文的记录直接丢弃；标题按领域关键词粗滤四领域，非目标领域与实习岗不入库。"),
])
H(2, "4.4  时滞、噪声与“通胀”问题的处理")
TABLE("招聘数据三类问题与系统对策", ["问题", "表现", "对策"], [
    ["时滞", "JD 长期挂在门户上不更新；能力要求已变但旧 JD 仍在", "双时间戳：观察时间与有效期分开；覆盖率按周期重算，本期不再出现的要求写 valid_to；岗位状态用 90 天窗口与持续月数计算，不看累计条数"],
    ["噪声", "福利、公司介绍中的技术词；通用素质（沟通、团队协作）；模型品牌裸提及", "切段后只允许从职责 / 要求段抽取；通用素质在抽取时标为 generic 并丢弃；品牌名默认只留在证据上下文，只有同句有 API 集成 / 部署 / 微调 / 评测等动作才映射为技能点"],
    ["抄袭", "同一模板被多家公司或多个渠道复制", "simhash 近重只留最早一条；独立源按规范化公司计票；同一份 JD 对同一技能点只投一张要求判定票"],
    ["通胀", "JD 把所有可能的技术都写成“必须”，导致岗位要求膨胀", "要求判定票：明确必备票或明确加分票需占已分类票≥60% 且来自≥2 个独立源，否则只生成审核提案；必备要求等价数上限 12、正式要求上限 24；新增必备超过 max(3, 上期×50%) 记为岗位要求异常并暂停该岗位诊断"],
], widths=[1.8, 5.6, 8.1], font_size=9.5)
H(2, "4.5  交叉验证与融合机制")
P("系统不把任何单一来源当作真相。一个技能点要成为岗位的正式要求，必须同时满足：出现在职责或要求段；在该岗位去重 JD 簇中的覆盖率≥30%；要求性质得到≥60% 的明确票且票来自≥2 个独立公司；抽取置信与独立源数量共同决定的置信层不为低。一个岗位要从候选进入萌芽，必须有≥3 个独立公司在 90 天窗口内发布相关 JD，且大模型判别其为新岗位而非别名。一条岗位定义声明要获批，必须有至少两个独立证据源。这些门槛全部是 pipeline/constants.py 中的可配置常量，改一处即可全局生效。")

# ================================================================== 5 新岗位发现
H(1, "5  新岗位发现与定义")
H(2, "5.1  发现流程与状态机")
P("新岗位发现与既有岗位能力更新共用同一条管线和同一道审核闸，只是出口不同。抽取得到的岗位名先与 17 个覆盖靶子（如大模型应用工程师、Agent 工程师、数据工程师、嵌入式智能工程师等规范岗位名）做嵌入对齐，余弦阈值 0.84（按 17 个靶子两两最高相似度 0.828 与别名变体 0.965 校准）。对齐成功则并入该岗位；对不上的 JD 以“标题 + 已抽出技能点名”做 bge 嵌入，用层次聚类形成最小 3 条的发现簇，再由大模型只对簇代表做三分类：新岗位 / 既有别名 / 噪声。别名写 ALIAS_OF 边并入既有岗位，噪声丢弃，新岗位进入状态机。")
FIGURE(f"{FIG}/fig_states.png", "新岗位发现流程与岗位状态机", 6.0)
TABLE("岗位状态与流转条件（pipeline/constants.py）", ["状态", "进入条件", "产品行为"], [
    ["候选 candidate", "聚类形成且判别为新岗位，但未达萌芽门槛", "未入谱；市场演化页可展示漏斗与卷宗；不可诊断、不可作为正式岗位；接口无口令时对候选岗 404 / 400"],
    ["萌芽 emerging", "近 90 天窗口内≥3 个独立源，且 LLM 判为新岗位而非别名", "入谱，标“新兴”；可进入图谱工作台；岗位定义获批且通过发布校验后可诊断"],
    ["成型 formed", "（≥10 独立源 或 持续≥6 个月）且岗位定义曾获批准或自动通过", "正式岗位；市场演化页按最近变化排序展示"],
], widths=[2.8, 6.0, 6.7], font_size=9.5)
P("状态由证据计票每日重算，不手写、不靠现场等满；萌芽与成型的个数随证据变化而变化，不是验收计数。")
H(2, "5.2  岗位定义生成与人工优化")
P("岗位定义是岗位在一个有效期内的正式说明，由大模型基于该岗位去重 JD 簇起草，拆分为可单独核对证据的“定义声明”，分为核心职责与典型行业应用场景两类；必备技能、加分技能不写在定义正文里，而由要求边（kind=required / bonus）与熟练级表达，避免定义与要求两处口径不一致。生成的岗位定义包含赛题要求的五个要素：")
TABLE("岗位定义要素与来源", ["要素", "来源与约束"], [
    ["岗位名称", "对齐到覆盖靶子的规范名，或簇判别产生的新名；别名通过 ALIAS_OF 并入"],
    ["核心职责", "定义声明（duty），每条声明至少两个独立证据源，逐条获批后整份定义才算获批"],
    ["必备技能", "kind=required 的有效要求边，附熟练级、独立源数、三类判定票与最短原文摘录"],
    ["加分技能", "kind=bonus 的有效要求边，同上"],
    ["典型行业应用场景", "定义声明（scenario），来自 JD 职责段中的业务场景描述，逐条挂证据"],
], widths=[3.4, 12.1], font_size=9.5)
P("人工优化在管理后台完成：新岗位首次发布作为审核提案进入待审队列，管理员可批准、改写后批准或驳回；提案原稿不可变，改写以审核决定的形式记录差异与理由；岗位定义未获批时前台显示“岗位定义审核中”，允许查看市场卷宗但不能进入简历诊断。每日管线以新证据重新计算覆盖率、票数与状态，形成动态更新。")
H(2, "5.3  示例：Agent 工程师")
P("赛题要求的新岗位示例为“Agent 工程师”。当前图谱中该岗位由 6 个独立源（小米、生数科技、深言、众擎机器人、智谱、小鹏汽车）的 JD 证据支撑，观察时间跨越 2025 年 4 月至 2026 年 8 月，因持续超过 6 个月且定义获批而处于成型状态；正式要求 22 项（必备 12、加分 10），观测中技能 175 项。要求边覆盖 AI Agent、Agent 系统、Function Calling、工具调用、Memory、LangChain、AutoGen、LLM、大模型、Python、Go、TypeScript、后端开发等技能点，每条都带 valid_from 与来源证据 ID。完整的岗位定义、要求边、证据与输入输出示例见提交物 data/eval-official-only/deliver/agent/（job.json、sources.jsonl、io.md）。")

# ================================================================== 6 既有岗位
H(1, "6  既有岗位能力动态更新")
H(2, "6.1  要求边生命周期与切片差分")
P("每条要求边带 valid_from 与 valid_to。字段（性质、熟练级、置信、来源）发生有效变化时不修改旧边，而是写旧边 valid_to 并新建一条边；本期不再出现的要求同样写 valid_to。切片差分即当前岗位在明确起止时间内相对上一周期的变化：新增（首次入池）、升值（覆盖率上升并刚过入池线）、失效（已写 valid_to）。产品默认比较自然季度，图谱工作台可切换周期并提供“近 90 天生效”筛选。演化事件按时间倒序列在市场演化页与首页流水，不做时间轴回放控件。")
H(2, "6.2  入池与要求判定票")
BUL([
    ("入池。", "技能点必须来自职责或要求段，且在岗位去重 JD 簇中的覆盖率≥30% 才写要求边；低于门槛记为观测中，只挂在岗位 watching 列表，求职者界面写“市场开始提，还没进要求，不算你的缺口”。"),
    ("要求判定票。", "每份去重 JD 对同一技能点只投一票：原文写“必须”“要求”“熟练掌握”或列入明确任职要求记明确必备；写“优先”“加分”记明确加分；普通提及记未标。未标票计入覆盖率与票数展示，但不进入性质判定分母。"),
    ("判定规则。", "明确必备票或明确加分票占已分类票≥60% 且来自≥2 个独立源才写相应 kind；两边都不满足时只生成审核提案。管理页同时展示三类票数与原文，人工决定不抹掉原票，下个周期按新证据重算。"),
])
H(2, "6.3  要求组与要求等价数")
P("同一份 JD 明确写“任选”“或”“至少两种”时，按原文建立要求组并读取 min_required；跨 JD 推断的可替代关系（如 Python、Java、Go 分别被不同公司要求）只有在支撑同一职责、单份 JD 内很少共同出现且组合覆盖率≥30% 时才生成要求组合并提案，人工批准前不成组。匹配与缺口集按组对账，不把组员各记一个缺口。发布数量按要求等价数计算：独立要求边计 1，要求组按 min_required 计数。")
H(2, "6.4  演化事件与更新说明")
P("既有岗位的每次能力变化都产出一条演化事件，携带发生时间、证据引用、置信度与审核状态，payload 中记录 skill_id、新旧边字段与三类票数。前台“岗位变化”明确标注新增、失效与修改的能力项，并附去重公司数与代表性原文摘录，满足赛题“提供更新说明及数据源，明确标注新增、删除、修改能力项”的要求。")
H(2, "6.5  人工优化：审核工作台")
P("管理后台按岗位与待发布版本分组，显示岗位定义、上期与本期要求等价数、变化摘要与异常原因；下方分新增要求、性质变化、技能合并提案、要求组合并提案与撤回证据。每项可单独批 / 改 / 驳，也可对岗位版本“一键全部批准”：批量批准跳过逐条阅读与独立审核模型，但不能跳过证据 ID 存在、证据未撤回、摘录存在、无重复有效要求和岗位定义非空的确定性校验；岗位要求异常只有填写放行理由后才能越过。审计记录保存共享管理员会话、时间、全部提案 ID 与放行理由。管理员还可撤回证据或事件，所有引用该证据的岗位事实自动重新计票。")
FIGURE(f"{SHOT}/admin-queue.png", "管理后台审核工作台（口令登录后）", 6.2)
H(2, "6.6  示例：大模型应用工程师与别名判别")
P("赛题要求的既有岗位示例为“大模型应用工程师”。管线从旧年快照与新年 JD 中计算覆盖率变化：跨过 30% 入池线的技能进入新要求边，不再出现的旧边写 valid_to，通用素质与无动作模型品牌以可回滚方式失效；“大模型应用开发工程师”被簇判别写为 ALIAS_OF，证明判别器能识别别名而不误报新岗。2026-09-02 的校准运行中，该岗位收敛为 12 条必备、23 条正式要求，14 条岗位定义声明均有至少两个独立证据源。输入输出示例见 data/eval-official-only/deliver/llm-app/（job.json、sources.jsonl、io.md、diagnose.example.json）。")

# ================================================================== 7 图谱
H(1, "7  新一代信息技术岗位全景图谱可视化")
H(2, "7.1  视图与颗粒度")
P("图谱工作台以三栏布局呈现：左栏为岗位列表与领域筛选，中栏为当前岗位的技能点切片画布，右栏为岗位说明与要求清单。颗粒度到技能点级别；技能点归入语言、框架、平台、工程、领域知识五个技能类目，即“按技术栈切换视图”；要求边上的 levels 属性支持按初级 / 中级 / 高级切换“级别视图”。工具栏还提供全部 / 必备 / 加分、近 90 天生效、类目、级别与关键词筛选，以及“拓扑 / 表格”双视图切换。")
FIGURE(f"{SHOT}/graph.png", "图谱工作台：Agent 工程师切片（拓扑视图）", 6.3)
H(2, "7.2  交互设计")
BUL([
    ("证据环布局。", "画布以岗位为中心，按证据强度与时间组织技能点：线越粗独立源越多，虚线为加分，悬停卡片显示原文摘录。力导向布局（d3-force）在 Web Worker 中迭代 300 步后稳定，防碰撞半径 140px 保证零重叠；可在“智能 / 固定”布局间切换。"),
    ("邻居高亮。", "点击节点以 BFS 计算邻接并高亮邻居、淡化其余节点与边；清除按钮恢复。"),
    ("判定摘要与证据抽屉。", "右栏每条正式要求先显示一句判定摘要（已分类票、明确性质票、独立源数），点“查看判断依据”从右侧滑出抽屉展示三类票数、簇内覆盖率、证据摘要与最短原文；关闭后焦点回到触发项。"),
    ("切片差分。", "岗位说明区固定展示获批岗位定义、典型职责、当前切片日期、上一周期、更新时间、数据状态以及本期新增 / 升值 / 失效计数，要求清单按必备、加分、观测中分组并另列变化清单。"),
    ("响应式与无障碍。", "桌面端默认图谱、移动端默认要求清单，切换不重置筛选；键盘可达、focus-visible、Esc 关闭抽屉、支持 320px 视口与 200% 缩放。"),
])
H(2, "7.3  首页总览与市场演化")
P("首页以“活数据总览”替代营销落地页：读数卡展示当前发布版本、入谱岗位数、成型 / 萌芽、候选中与去重 JD 样本数，实时流水显示最近的要求提案与拦截事件，岗位总览表按领域分组列出全部岗位的状态、独立源、本期增减与最近变化。市场演化页以求职者的问题组织岗位卷宗：为什么系统认为岗位正在形成、哪些公司开始招聘、本周期新增什么要求、与最接近的既有岗位有何区别、现在是否值得关注；候选 / 萌芽 / 成型作为筛选标签，观测中的技能以“审核中”计数展示。")
FIGURE(f"{SHOT}/home.png", "首页总览：发布版本、岗位状态与实时流水", 6.2)
FIGURE(f"{SHOT}/discover.png", "市场演化页：岗位卷宗与本周期变化", 6.2)

# ================================================================== 8 诊断
H(1, "8  人岗匹配度诊断与差距分析")
H(2, "8.1  诊断流程")
P("诊断在同一路由内依次完成上传简历、校对解析、选择岗位、生成对照、查看报告五步。全流程免登录，简历只保留在一小时有效的匿名会话中。")
FIGURE(f"{FIG}/fig_diagnose.png", "人岗匹配诊断流程", 6.3)
H(2, "8.2  简历解析")
BUL([
    ("文本提取。", "PDF 使用 pdfplumber 取文本层，docx 使用 python-docx；支持中英文混排与单栏 / 双栏；.doc、图片与扫描件直接返回 400 并提示不支持，不引入 OCR。原文件在解析完成或失败后立即删除。"),
    ("结构化抽取。", "大模型拆两个子任务并行：基本信息 + 教育 + 经历 + 项目；技能点列表（以当前图谱技能名作引导词表）。输出经 Pydantic 校验后全部过 align_skill 对齐到图谱技能点。"),
    ("简历证据片段与证据级。", "每个技能点必须能回到简历原文的最小摘录，证据级分三档：提及（只在技能栏或自我评价出现）、使用（经历或项目写明任务中的使用）、结果（写明本人责任以及数字、交付物或可核对结果）。证据级由确定性规则从句子中判断，不进入匹配分。"),
    ("用户校对。", "校对页要求确认当前角色、工作年限、学历、技能点、明确熟练级与证据级；解析修正必须有原文支持，用户补充但简历未写的技能单列为“你补充的，简历尚未证明”，只进入方向结论的补充说明。修正只写当前会话，不回写图谱、抽取器或评测数据。"),
])
FIGURE(f"{SHOT}/diagnose-step2.png", "诊断第二步：校对解析结果", 6.2)
H(2, "8.3  匹配算法与差距分析")
P("技能对齐采用统一入口 align_skill：先做确定性表面归一（大小写、全半角、空格、标点、已获批旧名称），再查获批同义词表，最后才用 bge 余弦（评测冻结阈值）。匹配分只在服务端计算：")
CODE("score = 100 × (必备覆盖 + 0.3 × 加分覆盖) / (必备满分 + 0.3 × 加分满分)\n\n必备未覆盖记 0；简历标了熟练级且低于岗位要求记 0.5（半档，进缺口集）；\n简历未标熟练级只比有无；加分缺失不伤必备分；边上 weight 不进分；分母为 0 则 score=0。\n档位：≥0.85 高度匹配 · ≥0.60 基本匹配 · ≥0.35 有明显差距 · 其余不匹配")
P("求职者看到的是档位文案而不是大数字。缺口集是目标岗位必备技能点中对齐后未覆盖或半档不足的集合，对外只写“简历未找到证据”，不推断求职者真实能力。换档条件 shift_set 对缺口与半档做最小补集计算：先列出单独补齐即可升档的技能点，再列出成对才能升档的组合，得到使档位升一档的最小技能集合，学习路径按此排序。要求组按 min_required 对账，不把组员各记一个缺口。")
H(2, "8.4  多维度匹配分析与方向结论")
P("系统对同一简历在当前发布版本的全部可诊断岗位中逐层排序：匹配档位 → 必备技能覆盖 → 有简历证据的岗位专属技能数 → 可迁移工程能力数，最后以数据新鲜度与证据充分度破同序，给出最多三个推荐岗位与每岗前两条理由；用户可改选两个岗位对照。方向结论由模型分别比较必备技能覆盖、简历证据级、可迁移工程能力、岗位专属经历与经验学历风险后解释“当前更接近哪个岗位”，不合成新的综合分；两个岗位在档位、必备覆盖、专属技能数、可迁移能力数和最小换档要求等价数全部相同时，系统明确写“当前简历证据不足以区分”，不强选胜者。")
FIGURE(f"{SHOT}/diagnose-step3b.png", "诊断第三步：岗位推荐序与双岗选择", 6.2)
H(2, "8.5  改进建议与学习路径规划")
TABLE("报告三视图与产品记忆点", ["视图 / 组件", "内容", "约束"], [
    ["结论", "方向结论、简历定位判断、邻近岗位迁移地图、两岗分维度比较、最多三项关键优势与三项主要风险、简历内容状态、档位与换档条件", "每条引用已有内容的判断可打开简历证据片段；缺失项标明检查了哪个简历部分，只写“简历中未找到”"],
    ["行动：双轨行动清单", "简历证明轨：分析全部经历与项目，默认展示影响最大的五条简历改写建议（原文、当前问题、建议版本、使用的原文事实、仍需补充的事实）；能力提升轨：按换档条件列最多三个任务，每项含技能点、岗位职责证据、学习资源、可写进简历的交付物", "原文没有的数字、责任、规模和结果只能写成待补占位；建议在重新上传前不能充当简历证据"],
    ["行动：学习路径", "按换档条件排序，默认最多 5 步；每步含技能点、职责段摘录、一条可打开的链接、为何先补它", "学习资源由 LLM 提出候选，必须通过 HTTP 访问校验且页面标题与技能点相符，缓存 Redis 7 天；不自建课程库"],
    ["行动：换档模拟器", "把缺口、熟练级不足或要求组候选技能暂时加入后即时重算两岗与邻近岗档位和换档条件", "结果固定标为“假设结果，尚未被简历证明”，不修改简历技能、证据级或正式结论"],
    ["行动：求职叙事稿", "面向推荐岗位的 45 至 60 秒自我介绍草稿", "只使用简历原文与用户确认的事实"],
    ["依据：简历证据地图", "左侧简历证据片段、右侧正式要求，连线标明提及 / 使用 / 结果", "未连接的要求写“简历未找到证据”，未连接的经历只写“与当前岗位要求未建立联系”"],
    ["依据：市场信号雷达", "按技能点列观测中的簇内覆盖率、独立源数、周期变化、代表性证据与未入正式要求原因", "固定说明“市场开始提，还没进要求，不算你的缺口”"],
    ["依据：对账表", "岗位关键词对照（已有证据 / 只有提及 / 未找到）、缺口集、熟练级差异、计分公式、岗位切片、对齐痕迹", "不显示 ATS 分数或命中率，半档 0.5 与加分 0.3 只写在这里"],
], widths=[3.2, 6.6, 5.7], font_size=9.5)
FIGURE(f"{SHOT}/diagnose-report.png", "诊断报告页：双岗对照（无法区分时分别给出换档条件）", 6.2)
FIGURE(f"{SHOT}/diagnose-report-detail3.png", "单岗详细报告：结论视图（档位、换档条件、简历分析与待核对证据）", 6.2)
H(2, "8.6  隐私与告知")
P("上传区在用户选择文件前写明支持格式、必要简历文本会发送给当前模型服务商、产品数据库不保存简历、匿名会话最长保留一小时；选岗页在生成前展示图谱发布版本、岗位更新时间、数据状态与独立源数。报告可复制一小时对照链接并显示准确失效时间，“打印 / 保存为 PDF”只调用浏览器打印能力，服务器不生成文件。产品分析不采集简历内容。")

# ================================================================== 9 幻觉防控
H(1, "9  幻觉防控与图谱构建科学性")
P("“能力幻觉”指模型把 JD 中不存在、不成立或不属于该岗位的能力写进图谱，或在报告中编造求职者不具备的经历。本系统把幻觉防控落实为贯穿抽取、归一、统计、审核、发布与报告六个层次的十道闸门，任何一道未过只会降级，不进入正式图谱。")
FIGURE(f"{FIG}/fig_gates.png", "图谱构建的幻觉防控闸门", 6.5)
H(2, "9.1  抽取层")
BUL([
    ("切段约束。", "JD 先切为职责、要求、福利、公司介绍四段，技能只允许从前两段抽出，福利与介绍中的词不入池。"),
    ("强 schema 与校验。", "模型必须输出 JSON，字段与枚举值（kind、proficiency、section、category、candidate_type、vote）全部由 Pydantic 校验，对不上即失败；失败整单重试，三次仍失败进待审并标记抽取失败。"),
    ("每项必带原文摘录。", "每个技能点携带包含技能名的最短原文片段（≤80 字）；无摘录的边视为无证据链，置信层强制为低。"),
    ("原文回指校验。", "预测技能必须能回指到原文词表命中，阻止语义嵌入“想象”出原文没有的技能；候选召回只接受精确原文词表命中。"),
    ("类型边界。", "通用素质（沟通、团队协作、学习能力）标为 generic 丢弃；GPT、Gemini 等模型品牌默认只留在证据上下文，同句有动作才映射为 API 集成、部署、微调、评测等技能点；CV、NLP 等宽泛领域词只有职责上下文充分时保留；“C/C++/Java”等并列表述必须拆开。"),
    ("提示词与技能定义单一权威。", "技能点定义常量 SKILL_DEFINITION 同时供金标起草与管线抽取使用，改定义即须按规程重新起草并裁决金标（ADR-0012）。"),
])
H(2, "9.2  归一层")
P("技能归一分三层：表面归一（大小写、全半角、空格、标点、已获批旧名称）直接对齐；跨语言、缩写与全称（如 Prompt Engineering 与提示词工程）只能生成技能合并提案，人工比较定义与原文后才写入共享同义词表并建立旧 ID 映射；嵌入相近但不相同的技术（LangChain 与 LangGraph、GPT 与 Gemini、RAG 与向量数据库）禁止合并，代码中维护禁止配对表。批准合并不删除旧技能 ID，历史要求边与证据继续可追溯。")
H(2, "9.3  统计层")
P("覆盖率、要求判定票、独立源与置信层四项统计把“模型说的”变成“市场证明的”。置信层按优先级：无证据链 → 低；≥3 独立源且抽取置信≥0.8 → 高；抽取置信≥0.5 → 中；其余 → 低。高层可由人工或自动审核批准；中层进入预览层等待人工审核，可在发现与管理场景查看但不进入正式图谱与诊断；低层永不自动入谱，管理员仍可批准但只记 approved 不记 auto_passed。")
H(2, "9.4  审核层")
BUL([
    ("待审队列是最后一道闸。", "新岗位首次发布、核心必备新增、低置信抽取、消解失败、技能合并提案、要求组合并提案共用同一队列；审核提案不可变，人工改写以决定形式记录差异。"),
    ("独立自动审核。", "自动审核开关默认关闭；开启后只让高置信事实依次接受确定性校验（有摘录、≥3 独立源、有 valid_from）与独立审核模型复核，抽取模型不能审核自己的输出（LLM_REVIEW_MODEL 必须与抽取模型不同），审核模型不可用或失败时 fail-closed。"),
    ("批量审核不越过完整性。", "批量批准可跳过逐条阅读，不能跳过证据存在、未撤回、摘录存在、无重复要求、定义非空校验。"),
])
H(2, "9.5  发布层")
BUL([
    ("诊断发布校验。", "每个可诊断岗位公开前接受确定性检查：定义非空、至少一组有效必备、每条要求有未撤回证据、同一技能点无重复有效要求、必备等价数≤12、正式要求等价数≤24、增量不超过 max(3, 上期必备×50%) 与 max(5, 上期正式×50%)。命中异常只暂停该岗位诊断，求职者看到“该岗位数据正在校验，暂不可诊断”，不会先看到污染后的结论。"),
    ("不可变发布版本。", "一次管线运行通过校验后原子发布一个图谱版本；公开指针可切到当前或历史版本；诊断会话固定使用创建时的版本。"),
    ("撤回与陈旧。", "证据撤回后所有引用它的事实重新计票，撤回不进入演化；连续 48 小时无成功批处理标记“数据陈旧”，页面继续展示上一成功版本并通知管理员。"),
])
H(2, "9.6  报告层")
P("诊断报告中的模型判断必须携带简历证据片段 ID 或岗位证据 ID；缺失判断携带被检查的简历部分，只写“简历中未找到”；简历改写建议只能重组原文事实，原文没有的数字、责任、规模和结果只能写成待补占位；学习资源必须通过 HTTP 访问与标题校验；方向结论不合成分数，无法区分时明确说明。这样即便模型出现幻觉，也无法以“事实”的形式进入用户视野。")

# ================================================================== 10 大模型 / 图谱 / RAG
H(1, "10  大模型、知识图谱与检索增强的协同")
TABLE("大模型在系统中的职责边界", ["环节", "模型任务", "确定性约束"], [
    ["JD 抽取", "从职责 / 要求段输出技能点 JSON（名称、性质票、熟练级、置信、摘录、段落、类目、候选类型）", "Pydantic 校验、原文回指、切段、类型边界、最多 40 条"],
    ["簇判别", "对发现簇代表三分类：新岗位 / 别名 / 噪声", "输入仅簇代表与技能名；状态流转由独立源计票决定，模型不定状态"],
    ["岗位定义起草", "起草核心职责与典型场景的定义声明", "每条声明需≥2 独立证据源，人工逐条批准"],
    ["简历解析", "两个并行子任务输出结构化 JSON", "技能点全部过 align_skill；证据级由规则判定；用户校对"],
    ["方向结论解释", "基于确定性比较结果解释更接近哪个岗位、优势与风险", "输入为图谱要求边 + 简历证据片段（检索增强），每条判断带证据 ID，不合成分数"],
    ["简历改写与叙事稿", "重组原文事实生成建议版本与 45 至 60 秒叙事", "原文没有的事实只写待补占位"],
    ["学习资源", "给出技能点对应的候选学习页面 URL", "HTTP 可访问且标题匹配才发布，缓存 7 天"],
    ["金标起草", "只读原文起草金标草稿", "人工两段裁决，禁止以预测为唯一依据改金标（ADR-0011）"],
], widths=[2.8, 6.4, 6.3], font_size=9.5)
P("知识图谱是模型输出的“落点”与“边界”：模型只能把技能点落到已存在或经审核新建的 Skill 节点上，只能给 Job 与 Skill 之间的 REQUIRES 边投票，而不能凭空创建岗位事实。诊断解释采用检索增强方式：先从图谱发布版本中检索目标岗位的有效要求边、证据摘录与观测中技能，从简历会话中检索证据片段，再把这些结构化事实作为唯一上下文交给模型生成解释，模型的每句话都必须回指上下文中的证据 ID。嵌入模型 bge-m3 承担技能对齐、实体消解与岗位聚类三项任务，阈值均按金标校准并在评测时冻结。")
P("所有生成调用经过唯一出口 complete_json(messages) -> dict，抽取与解析使用非思考模式和 JSON 模式以降低延迟；超时 60 秒，失败重试一次，若首个 JSON 截断则第二次追加紧凑输出约束并将输出上限减半；外部模型设置每日调用量与费用硬上限，额度耗尽时明确失败而不静默切换供应商。")

# ================================================================== 11 实现
H(1, "11  系统实现")
H(2, "11.1  代码结构与规模")
CODE("apps/api/                 FastAPI 后端（约 9,000 行 Python）\n  app/main.py             路由：公开只读、诊断、管理、SSE\n  app/graph.py            Neo4j Cypher 封装：岗位 / 技能 / 要求边 / 事件 / 发布版本 / 审核决定\n  app/llm/client.py       多供应商 chat 唯一出口   app/llm/embed.py  bge-m3 嵌入\n  app/collectors/         ats.py 门户适配 · controller.py 去重落盘 · simhash.py · normalize.py · sink.py\n  app/pipeline/           sections 切段 · extract 抽取 · align 归一 · gate 入池/判定/置信/审核 · discover 簇判别\n                          status 状态机 · diagnostic_release 发布校验 · curate_public 公开校准 · constants\n  app/matching/           resume 解析 · score 匹配 · bands 档位/换档 · report 报告 · session 会话\n  app/eval/               f1 · run 三项评测 · draft 金标起草 · adjudicate 裁决 · deliver 两岗提交物\napps/web/                 Next.js 前端（约 6,000 行 TS/TSX）\n  app/page.tsx home.tsx   首页总览      app/graph/    图谱工作台（flow-canvas、workbench）\n  app/diagnose/           五步诊断      app/discover/ 市场演化    app/admin/  审核 / 裁决 / 发布 / 门户\ntests/                    22 个测试文件，180 个用例\ndata/official-only/jd/    官方 JD 快照（3,595 条）      data/eval-official-only/  金标与提交物\ndata/snapshot/            图谱发布快照（reviewed.json / release-*.json）\ndocs/                     product.md · tech.md · verification.md · frontend.md · 55 篇 ADR · 调研\ndocker-compose.yml · .github/workflows/ci.yml · CONTEXT.md（领域术语）")
H(2, "11.2  后端接口")
TABLE("主要 API 接口", ["方法", "路径", "用途"], [
    ["GET", "/meta、/v1/meta", "四领域、当前图谱发布版本、模型供应商、会话保留时长、运行状态、数据陈旧标记"],
    ["GET", "/jobs、/jobs/{id}", "岗位列表与详情（获批定义、典型职责、独立源数、更新时间、数据状态）；无口令时不返回候选岗"],
    ["GET", "/graph/jobs/{id}", "当前与上一周期切片、类目、按必备 / 加分 / 观测中分组的技能点、判定摘要与票数、period_delta"],
    ["GET", "/discover、/discover/{id}", "候选 / 萌芽 / 成型看板与岗位卷宗"],
    ["GET", "/feed、/pulse", "故事、计数、管线、热度、演化流水"],
    ["POST", "/sessions；PUT /sessions/{id}/skills", "上传简历创建会话；会话内修正技能点与熟练级"],
    ["POST", "/diagnose/recommend", "岗位推荐序（最多三个可诊断岗位与理由）"],
    ["POST", "/diagnose", "单岗或双岗对照报告（方向结论、缺口集、换档条件、学习路径、证据地图、迁移地图、市场雷达等）"],
    ["POST", "/diagnose/simulate", "换档模拟：假设技能加入后的档位与换档条件，不写会话"],
    ["POST/DELETE", "/admin/login、/admin/logout", "共享口令常量时间比较，签发短期管理会话 Cookie"],
    ["GET/POST", "/admin/queue、/admin/queue/{id}/approve|reject", "待审队列与逐条审核（可带改写 payload）"],
    ["POST", "/admin/jobs/{job}/versions/{ver}/approve-all", "岗位版本批量批准，异常需 override_reason"],
    ["GET/POST", "/admin/diagnostic-release、/admin/public-curation", "诊断发布校验结果与公开校准"],
    ["GET/PUT", "/admin/passthrough", "自动审核开关"],
    ["POST", "/admin/releases、/admin/releases/{id}/rollback", "发布图谱版本与回滚"],
    ["POST", "/admin/events/{id}/retract、/admin/evidence/{id}/retract", "撤回事件 / 证据"],
    ["GET/POST", "/admin/adjudicate/next、/admin/adjudicate/decide", "金标裁决工作台"],
    ["GET/POST", "/admin/portals、/admin/collect、/admin/ops/status、/admin/review/stats", "门户管理、触发采集、运行状态、审核统计"],
    ["GET", "/events/stream", "SSE 事件流（管理会话），支持 Last-Event-ID 续传"],
], widths=[1.8, 5.6, 8.1], font_size=9)
H(2, "11.3  前端实现")
P("前端采用“纸面工作台”设计系统：IBM Plex 字体、直角、发丝分割线、单一强调色，所有颜色与字体只引用设计令牌；动效只动 transform 与 opacity。主导航为首页、图谱工作台、市场演化、管理后台与“开始诊断”主按钮。图谱画布基于 React Flow，力导向布局在 layout.worker.ts 中以 d3-force 计算；诊断页为受控向导，五步状态保存在组件内并支持对照链接直达；管理后台包含审核队列、岗位版本发布、金标裁决与门户管理四个看板。构建通过 npm run typecheck 与 npm run build，图谱页首屏 JS 约 174 KB。")
H(2, "11.4  部署")
FIGURE(f"{FIG}/fig_deploy.png", "容器化部署拓扑", 6.0)
P("部署只需三步：复制 .env.example 为 .env 并填写模型供应商密钥与管理口令；执行 docker compose up -d --build；访问前端 3000 端口、API 8000 端口与 Neo4j 7474 控制台。首次启动会构建 api 与 web 镜像并启动每日管线容器。图谱数据有两种填充方式：默认由每日管线采集、抽取并发布（需模型密钥与门户访问）；演示或冷启动时在 .env 中设置 SNAPSHOT_PATH=/app/data/snapshot/reviewed.json，API 启动检测到空图会自动导入 data/snapshot/reviewed.json（与 release-2026-09-05.json 字节一致，2026-09-05 由产品图以 --slim 方式导出，含公开 release 预计算快照、审核提案与决定），已在临时空 Neo4j 卷上验证恢复后 /meta、/jobs、/discover、/graph/jobs/{id} 均正常。")
CODE("cp .env.example .env            # LLM_PROVIDER=deepseek  DEEPSEEK_API_KEY=...  ADMIN_PASSWORD=...\ndocker compose up -d --build     # web :3000  api :8000  neo4j :7474/7687  redis :6379  pipeline\ndocker compose ps                # 健康检查：/meta、redis-cli ping、Neo4j 7474\n# 冷启动可选：.env 追加 SNAPSHOT_PATH=/app/data/snapshot/reviewed.json 后再 up，空图自动导入\n\n# 手动运行管线与评测\nPYTHONPATH=apps/api python -m app.collectors --daily\nPYTHONPATH=apps/api python -m app.pipeline\nEVAL_DIR=data/eval-official-only PYTHONPATH=apps/api python -m app.eval report")
TABLE("主要环境变量", ["变量", "作用"], [
    ["LLM_PROVIDER / *_API_KEY / *_MODEL", "选择 deepseek、bai 或 tuzi 供应商及对应密钥与模型"],
    ["EMBED_API_KEY / EMBED_MODEL", "硅基流动 bge-m3 嵌入；未设置时回落本地哈希向量（仅测试）"],
    ["ADMIN_PASSWORD", "管理后台共享口令"],
    ["NEO4J_URI / NEO4J_USER / NEO4J_PASSWORD", "图数据库连接"],
    ["REDIS_URL", "会话、事件流与指纹集合"],
    ["DATA_DIR / JD_DIR / EVAL_DIR", "JD 快照与评测金标目录"],
    ["CORS_ORIGINS / NEXT_PUBLIC_API_URL", "允许的前端来源与前端访问的 API 地址"],
    ["ALIGN_THRESHOLD / EXTRACT_WORKERS / EVAL_WORKERS", "对齐阈值（评测冻结）、抽取与评测并发"],
], widths=[6.5, 9.0], font_size=9.5)
H(2, "11.5  运行保障")
BUL([
    "每日管线独立容器运行，采集后串行抽取，失败样本按检查点重试；连续 48 小时无成功批处理标记数据陈旧。",
    "结构化日志记录请求 ID、路由、状态码、耗时、模型、Token、费用与错误类型，保留 14 天，不记录简历正文、口令与完整会话 ID。",
    "图谱以版本化 JSON 快照备份，恢复只走幂等导入器；管线失败、数据陈旧、备份失败与发布失败统一发送到可配置 Webhook，管理页显示最近运行状态。",
    "测试库 neo4j-test 使用独立数据卷与端口，pytest 与 CI 永不写产品图。",
])

# ================================================================== 12 测试
H(1, "12  测试方案与验证")
H(2, "12.1  测试数据集")
P("评测数据全部来自官方门户 JD 快照与自建简历集，存放于 data/eval-official-only/，JSONL 一行一条，技能点一律写图谱 Skill.id，原文别名放 mentions。对齐阈值、模型名与评测日期冻结在 freeze.json，评测只读该文件。")
TABLE("测试数据集规模与构成", ["数据集", "规模", "构成与要求"], [
    ["JD 金标 jd.jsonl", "100 条（去重后）", "四领域均有，人工智能不少于 40 条、其余领域各不少于 10 条；每条标 job_id、skills[]（id、kind、proficiency）、mentions[]、section、watching[]"],
    ["简历金标 resume.jsonl", "100 份", "覆盖单栏 / 双栏、应届 / 社招、PDF / docx；标技能点集合与姓名 / 教育 / 经历；双栏样本打 layout: split 标签"],
    ["匹配金标 match_pairs.jsonl", "100 对", "金标简历 × 目标岗位的交叉抽样，人工标缺口集 gap_ids；半档只在简历带熟练级且低于岗位时计入"],
    ["两岗提交物 deliver/", "Agent 工程师、大模型应用工程师", "job.json（定义、状态、独立源、要求边）、sources.jsonl（≥3 独立源证据）、io.md（输入 JD 摘录 → 输出技能点、入谱结果、演化事件）、diagnose.example.json、dual-diagnose.redacted.json"],
    ["JD 快照 data/official-only/jd/", "3,595 条", "管线原始输入，每条含公司、岗位名、正文、观察时间、渠道、来源 URL、simhash"],
], widths=[4.0, 3.2, 8.3], font_size=9.5)
P("金标建立遵循“LLM 起草、人工裁决”规程（ADR-0011）：模型只读原文起草草稿并留痕；修订分两段，第一段盲改只看原文与技能词表，第二段允许用系统预测找分歧，但结论只认原文证据并把理由写进 notes，禁止以预测为唯一依据改金标；修订后必须重跑未 mock 的三项评测。")
H(2, "12.2  指标口径")
P("三项准确率均为技能点集合的 set-based F1：预测集与金标集先过 align_skill（读 freeze.json），再计算精确率、召回率与 F1，逐条求平均；空预测且空金标记 F1=1，一边空一边非空记 0。")
TABLE("三项准确率定义", ["指标", "预测", "金标", "达标"], [
    ["JD 解析准确率", "管线从 JD 抽出并对齐的技能点 id 集合", "jd.jsonl 中 skills[].id", "F1 ≥ 0.90，≥100 条"],
    ["简历提取准确率", "管线从简历抽出并对齐的技能点 id 集合", "resume.jsonl 中 skills[].id", "F1 ≥ 0.90，100 份"],
    ["人岗匹配准确率", "系统对（金标简历技能，金标岗位要求边）算出的缺口集", "match_pairs.jsonl 中 gap_ids", "F1 ≥ 0.90，≥100 对；不喂解析输出"],
], widths=[3.2, 5.0, 4.0, 3.3], font_size=9.5)
P("学习路径不定准确率，改为抽检换档条件上的技能点是否都有一条可打开的学习资源链接；档位一致率与端到端 PDF→缺口 冒烟测试另报，不进达标口径。")
H(2, "12.3  评测流程与自动化")
CODE("docker compose --profile test up -d neo4j-test      # 独立测试库\nEVAL_DIR=data/eval-official-only PYTHONPATH=apps/api python -m app.eval jd\nEVAL_DIR=data/eval-official-only PYTHONPATH=apps/api python -m app.eval resume\nEVAL_DIR=data/eval-official-only PYTHONPATH=apps/api python -m app.eval match\nEVAL_DIR=data/eval-official-only PYTHONPATH=apps/api python -m app.eval report  # 汇总\nPYTHONPATH=apps/api TEST_NEO4J_URI=bolt://localhost:17687 pytest --cov -q  # 门禁 60%")
P("任一 F1 低于 0.90 时评测脚本退出码为 1。GitHub Actions 在每次推送时以 mock 模型运行三项评测与全部单元测试并执行覆盖率门禁；正式报告中的数字只来自未 mock 的本地运行，summary.md 记录三项 F1、覆盖率、学习路径抽检覆盖与 freeze.json 哈希。改动提示词或 freeze.json 必须重跑未 mock 的三项。")
H(2, "12.4  当前验证结果")
TABLE("最近一次完整验证结果", ["项目", "结果", "口径 / 备注"], [
    ["JD 解析 F1", "0.930", "2026-09-05 未 mock 完整运行，100 条；达标，差距样本与复跑记录于 summary.md，未修改金标"],
    ["简历提取 F1", "1.000", "同上，100 份"],
    ["人岗匹配（缺口集）F1", "1.000", "同上，100 对，喂金标不喂解析输出"],
    ["单元测试", "179 passed，1 skipped", "2026-09-05 在干净 python:3.12 容器内运行，22 个测试文件，连接独立测试库"],
    ["行覆盖率", "69.67%", "pytest-cov，分母排除 collectors 与 client.py 网络分支；门禁 60%"],
    ["前端", "typecheck 通过，build 通过", "Next.js 15 生产构建"],
    ["Compose", "五容器 healthy", "docker compose ps；2026-09-05 空卷导入 data/snapshot/release-2026-09-05.json 成功，API 公开接口正常"],
    ["浏览器验收", "首页、图谱、发现、诊断、管理均 200，无 4xx/5xx", "系统 Chrome + Playwright；PDF 与 docx 上传完成单岗 / 双岗报告；.doc 返回 400；320px 与 200% 缩放无溢出；键盘 Esc、focus-visible、reduced-motion 复核"],
], widths=[3.6, 3.6, 8.3], font_size=9.5)
P("JD 解析 F1 从 2026-09-02 的 0.814 提升到 0.930，主要改进为：单段抽取取代两段式（两段式实测召回从 0.729 降至 0.604 后回退）、原文词表候选召回只接受精确命中并复用生产管线、评测并发降至 2 至 4 路避免长 JSON 截断。剩余差距集中在领域知识类技能的召回，后续在不改金标的前提下继续调整提示词与召回策略，并按规程重跑。")
H(2, "12.5  单元测试")
TABLE("单元测试覆盖重点与模块覆盖率", ["模块", "覆盖率", "重点用例"], [
    ["pipeline/extract.py", "94%", "JSON 字段归一、枚举容错、要求判定票、并列拆分、通用素质与品牌边界、原文词表召回"],
    ["pipeline/gate.py", "81%", "置信层四种分支、覆盖率入池与福利段不计、判定票 60% / 2 源、直通开关、低置信不可 auto_passed、抽取失败入队、事件应用"],
    ["pipeline/sections.py / constants.py", "100%", "四段切分与摘录归段"],
    ["pipeline/status.py", "88%", "3 源 / 90 天萌芽；10 源或 6 个月且定义获批成型；渠道与近重不计票"],
    ["pipeline/diagnostic_release.py", "93%", "定义缺失、必备缺失、证据缺失 / 撤回、重复要求、数量上限、增量异常与放行"],
    ["pipeline/curate_public.py", "89%", "公开校准与要求排序"],
    ["pipeline/align.py", "71%", "表面归一、同义词命中、余弦过线与未命中、禁止配对、复合拆分、聚类"],
    ["matching/bands.py / score.py", "97% / 98%", "四档阈值、半档 0.5、加分缺失不伤必备、分母为 0、换档条件单独项优先于成对项"],
    ["matching/resume.py", "85%", "PDF / docx 文本、证据级三档、日期冲突、技能对齐、不支持格式 400"],
    ["matching/report.py", "68%", "推荐序、方向结论与无法区分、证据地图、换档模拟、迁移地图、市场雷达、资源 URL 校验"],
    ["graph.py", "77%", "版本化要求、发布与回滚、审核决定与批量决定、撤回、别名、看板与卷宗"],
    ["eval/f1.py / run.py / schema", "—", "F1 函数空集规则、JSONL 必填字段校验、mock 评测流程"],
    ["main.py（路由）", "49%", "公开只读 200 / 404、候选岗位 400、会话隔离、管理登录 401、批量审核幂等、SSE"],
], widths=[4.6, 1.8, 9.1], font_size=9.5)
H(2, "12.6  功能测试用例（节选）")
TABLE("功能测试用例节选", ["编号", "场景", "输入", "预期结果"], [
    ["TC-01", "JD 去重", "同一 JD 由两个门户重复采集", "只落一条快照，第二条命中 ingest:fp 跳过"],
    ["TC-02", "近重去重", "两份正文 Hamming 距离 ≤3 的 JD", "只保留观察时间最早的一条计入独立源"],
    ["TC-03", "福利段噪声", "福利段出现“Python 培训”", "不产生技能点、不计覆盖率"],
    ["TC-04", "通用素质", "要求段“良好的沟通能力”", "candidate_type=generic，丢弃"],
    ["TC-05", "品牌边界", "“基于 GPT 进行微调”", "生成“微调”技能点，GPT 留在摘录"],
    ["TC-06", "并列拆分", "“熟悉 C/C++/Java”", "拆为 C、C++、Java 三条，摘录相同"],
    ["TC-07", "判定票不足", "3 票中 1 票明确必备、2 票未标", "不写正式要求边，只生成审核提案"],
    ["TC-08", "萌芽流转", "近 90 天 3 家公司 JD，LLM 判为新岗位", "状态 candidate → emerging"],
    ["TC-09", "别名判别", "“大模型应用开发工程师”簇", "写 ALIAS_OF → 大模型应用工程师，不占候选列"],
    ["TC-10", "低置信直通", "layer=low 且直通开关开启", "仍进待审，不记 auto_passed"],
    ["TC-11", "发布校验", "某岗位必备等价数 13", "该岗位排除出推荐与诊断，其他岗位正常"],
    ["TC-12", "批量审核", "岗位版本含撤回证据的提案", "整批拒绝，不写任何决定"],
    ["TC-13", "简历格式", "上传 .doc 或扫描件", "HTTP 400，中文提示不支持"],
    ["TC-14", "半档", "简历 Python=aware，岗位要求 able", "覆盖记 0.5，进入缺口集与换档条件"],
    ["TC-15", "无法区分", "两岗档位、覆盖、专属技能、可迁移能力、换档等价数全相同", "报告写“证据不足以区分”，分别给换档方案"],
    ["TC-16", "换档模拟", "把缺口技能加入假设集", "返回假设档位与新换档条件，会话技能不变"],
    ["TC-17", "会话过期", "对照链接 session 已过期", "回到 idle 并提示重新上传"],
    ["TC-18", "管理登录", "错误口令", "401，页面显示“口令错误”，登录限速"],
], widths=[1.4, 2.6, 5.6, 5.9], font_size=9)

# ================================================================== 13 创新
H(1, "13  技术创新点")
NUM([
    ("双时间演化图谱。", "观察时间与有效期分离，旧边写失效不删，切片差分替代静态画像，天然解决 JD 时滞并支持任意时点重演。"),
    ("多源交叉验证融合机制。", "官方门户直采 + fingerprint 幂等 + simhash 近重 + 规范化公司独立源计票 + 覆盖率入池 + 三态要求判定票 + 要求等价数上限，系统性解决噪声、抄袭与通胀。"),
    ("十道闸门的幻觉防控。", "抽取强 schema 与原文回指、三层归一禁止自动语义合并、置信层、抽取模型不得自审的独立审核、诊断发布校验与不可变发布版本，任何未过闸事实只降级不入谱。"),
    ("确定性优先的人岗诊断。", "匹配分、档位、缺口集、换档条件与岗位推荐序全部可复现，模型只做检索增强的解释且每句带证据 ID；引入简历证据级、要求组对账、换档模拟、无法区分判定，避免“AI 综合评分”式的伪精确。"),
    ("证据驱动的岗位状态机。", "候选 / 萌芽 / 成型由独立源计票与窗口自动流转，别名以 ALIAS_OF 并入，不手写状态，萌芽与成型个数随证据变化。"),
    ("可审计的人工优化闭环。", "提案不可变、决定留痕、批量审核不越过完整性、撤回重新计票、金标两段裁决，人工与模型的每一步都可追溯。"),
])

# ================================================================== 14 实用价值
H(1, "14  实用价值与可迁移性")
BUL([
    ("对企业。", "HR 与业务负责人可以在管理后台看到每个岗位要求的市场覆盖率、独立公司数与原文依据，用数据校准自家 JD，避免“通胀”式招聘要求；新岗位漏斗帮助企业更早识别新兴岗位的能力结构。"),
    ("对求职者。", "免登录上传简历即可获得基于证据的方向结论、缺口集、换档条件、双轨行动清单与学习路径，回答“卡在哪、换邻近岗会不会更好、这个月先补哪几样”，而不是一个含义模糊的匹配分。"),
    ("对高校与培训机构。", "全景图谱到技能点级别的覆盖率与周期变化，可直接用于课程设置与培养方案的动态调整。"),
    ("可迁移性。", "四领域框架与 17 个覆盖靶子只是对齐名单，岗位、技能点、要求边与状态全部由管线从数据长出；迁移到其他新一代信息技术岗位或其他行业只需新增采集门户与对齐靶子，抽取、归一、计票、审核、发布与诊断逻辑无需改动。模型供应商与嵌入服务均可替换，常量集中于 constants.py 与 bands.py。"),
])

# ================================================================== 15 提交物
H(1, "15  作品提交物清单")
TABLE("提交物与赛题要求对照", ["赛题要求", "提交物", "位置"], [
    ["作品设计实现方案", "本文档", "根目录 .docx"],
    ["PPT 作品介绍", "作品介绍演示文稿", "随压缩包提交"],
    ["10 分钟以内演示视频", "含新岗位（Agent 工程师）发现与既有岗位（大模型应用工程师）能力更新的图谱演示", "随压缩包提交"],
    ["源代码", "开源仓库（私有仓库开放评审权限）", "GitHub kangvcar/JobEvolution"],
    ["部署说明", "README.md、docker-compose.yml、apps/api/Dockerfile、apps/web/Dockerfile、.env.example", "仓库根目录与 apps/"],
    ["单元测试用例（覆盖率≥60%）", "22 个测试文件、180 个用例，覆盖率 69.67%，.coveragerc 门禁 60%，CI 工作流", "tests/、.coveragerc、.github/workflows/ci.yml"],
    ["测试数据：1 个新岗位与 1 个既有岗位的能力图谱及数据源（含输入输出示例）", "deliver/agent/ 与 deliver/llm-app/（job.json、sources.jsonl、io.md、diagnose.example.json）", "data/eval-official-only/deliver/"],
    ["≥100 条岗位 JD 及测试用例", "jd.jsonl（100）、resume.jsonl（100）、match_pairs.jsonl（100）、freeze.json、skills.json；官方 JD 快照 3,595 条", "data/eval-official-only/、data/official-only/jd/"],
], widths=[4.4, 6.6, 4.5], font_size=9.5)

# ================================================================== 附录
H(1, "附录 A  核心术语表")
TABLE("核心术语（摘自 CONTEXT.md）", ["术语", "定义"], [
    ["岗位 / 岗位定义 / 定义声明", "岗位是职位的规范化表示；岗位定义是有效期内的正式说明；定义声明是可单独核对证据的最小陈述"],
    ["可诊断岗位 / 诊断发布校验", "同时具备获批定义、有效必备要求与完整证据链并通过确定性校验的岗位；校验拦截空定义、无证据、重复要求、撤回证据与要求异常"],
    ["技能点 / 技能类目 / 技术品牌表述", "技能点是可由原文证明的原子技术技能；类目是导航层不参与对账；品牌名默认只是证据上下文"],
    ["要求边 / 要求判定票 / 要求组 / 要求等价数", "岗位指向技能点的核心边；每份 JD 一票三态；可替代技能组带 min_required；发布计数口径"],
    ["独立源 / 簇内覆盖率 / 观测中", "规范化公司名去重计票；技能在岗位 JD 簇中出现比例，≥30% 入池；未达门槛的技能"],
    ["双时间 / 演化事件 / 切片差分 / 撤回", "观察时间与有效期分开；变更原子记录；周期间要求边变化；作废不成立事实，不算演化"],
    ["岗位状态", "候选 → 萌芽（≥3 源、90 天、判新）→ 成型（≥10 源或≥6 月，且定义获批）"],
    ["置信层 / 待审队列 / 自动审核开关 / 审核提案与决定", "高 / 中 / 低三级；入谱前的人工闸；默认关闭的独立模型复核；不可变原稿与留痕决定"],
    ["匹配分 / 缺口集 / 换档条件 / 换档模拟", "内部技能覆盖分；必备未覆盖集合；升一档的最小补集；假设试算不改证据"],
    ["简历证据片段 / 简历证据级 / 解析修正 / 用户补充技能", "支撑技能点的原文摘录；提及 / 使用 / 结果三档；有原文支持的修正；无原文支持的声明"],
    ["方向结论 / 无法区分 / 岗位推荐序 / 双轨行动清单", "两岗比较结论不合成分数；五项全同时不强选；逐层排序不合成推荐分；简历证明轨 + 能力提升轨"],
    ["三项准确率 / 金标修订", "技能点集合 set-based F1 ≥0.90；盲改后裁决、只认原文证据、修订后重跑"],
], widths=[5.0, 10.5], font_size=9.5)

H(1, "附录 B  架构决策记录（ADR）索引")
P("仓库 docs/adr/ 下共 55 篇架构决策记录，记录了关键设计取舍及其理由，节选如下：")
TABLE("关键 ADR 节选", ["编号", "决策"], [
    ["0001", "覆盖靶子是对齐名单而非验收计数"],
    ["0002 / 0003", "产品图不是测试库；图谱可从快照重建"],
    ["0004", "匹配金标独立于解析输出"],
    ["0011", "LLM 起草金标，人工裁决"],
    ["0012", "技能点定义单一权威出处"],
    ["0016 / 0028 / 0031", "只发布版本化获批事实；版本化事实是节点；图谱发布不可变"],
    ["0017", "独立的自动审核模型，抽取模型不能自审"],
    ["0018 / 0032 / 0034", "匿名简历一小时；解析临时且告知；产品分析不含简历内容"],
    ["0022", "撤回不是演化"],
    ["0023 / 0050", "可替代要求成组；跨 JD 要求组需人工审核"],
    ["0025 / 0053", "审核提案不可变；批量审核跳过阅读不跳过完整性"],
    ["0033 / 0051", "只有合格岗位允许诊断；岗位要求异常暂停诊断"],
    ["0041 / 0042 / 0043", "方向结论无综合分；简历证据三级；改写不得虚构事实"],
    ["0047 / 0048 / 0049", "品牌是证据先于技能；技能归一三层；要求性质需明确票"],
    ["0052", "换档模拟不改变证据"],
    ["0054 / 0055", "JD 快照 ID 含正文哈希；每日 worker 先采集后抽取"],
], widths=[3.0, 12.5], font_size=9.5)

doc.save(OUT)
print("saved", OUT)
