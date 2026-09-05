# -*- coding: utf-8 -*-
"""Build 系统部署说明 (Word) for XH-202621 智演 JobEvolution.

解析 docs/submission/部署说明.md，输出同名 .docx。样式与 build_doc.py 一致：
正文宋体、标题黑体、西文 Times New Roman、代码块 Consolas + 灰底。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("/Users/kangvcar/Documents/code/JobEvolution/docs/submission/部署说明.md")
OUT = Path("/Users/kangvcar/Documents/code/JobEvolution/部署说明.md.docx")
# 项目里同类的参赛文档用 .docx 命名；部署说明随提交物命名，去掉标题中的间隔号避免落库乱码
OUT = Path("/Users/kangvcar/Documents/code/JobEvolution/智演JobEvolution_部署说明.docx")

BODY_FONT_CN = "宋体"
HEAD_FONT_CN = "黑体"
BODY_FONT_EN = "Times New Roman"
CODE_FONT = "Consolas"

doc = Document()

# ---------------------------------------------------------------- page setup
for s in doc.sections:
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.6)

settings = doc.settings.element
upd = OxmlElement("w:updateFields")
upd.set(qn("w:val"), "true")
settings.append(upd)

TAB_NO = [0]


def set_run_font(run, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=None, bold=None, color=None):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def style_font(style, cn, en, size, bold=False, color=(0, 0, 0)):
    style.font.name = en
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)


normal = doc.styles["Normal"]
style_font(normal, BODY_FONT_CN, BODY_FONT_EN, 12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(4)

for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12.5)):
    st = doc.styles[name]
    style_font(st, HEAD_FONT_CN, BODY_FONT_EN, size, bold=True)
    st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(6)


def H(level, text):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, cn=HEAD_FONT_CN, bold=True, color=(0, 0, 0))
    return p


def P(text, bold=False, indent=True, size=None, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def PR(parts, indent=True):
    """paragraph with mixed bold runs: list of (text, bold)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    for text, bold in parts:
        r = p.add_run(text)
        set_run_font(r, bold=bold)
    return p


def BUL(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            head, rest = it
            r = p.add_run(head)
            set_run_font(r, bold=True)
            r2 = p.add_run(rest)
            set_run_font(r2)
        else:
            r = p.add_run(it)
            set_run_font(r)


def NUM(items):
    BUL(items, style="List Number")


def CODE(text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.name = CODE_FONT
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "微软雅黑")
        rfonts.set(qn("w:ascii"), CODE_FONT)
        rfonts.set(qn("w:hAnsi"), CODE_FONT)
        r.font.size = Pt(9.5)
        ppr = p._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F3F3")
        ppr.append(shd)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)


def shade(cell, fill):
    tcpr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def TABLE(header, rows, font_size=10.5):
    TAB_NO[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"表 {TAB_NO[0]}")
    set_run_font(r, cn=HEAD_FONT_CN, size=10.5, bold=True)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        rr = p.add_run(h)
        set_run_font(rr, cn=HEAD_FONT_CN, size=font_size, bold=True)
        shade(c, "E7E6E6")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(str(v))
            set_run_font(rr, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ---------------------------------------------------------------- markdown parse

def strip_inline(text: str) -> str:
    """去掉行内 Markdown 标记（** **、` `、* * 等），保留纯文本。"""
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)  # 链接
    text = re.sub(r"`([^`]+)`", r"\1", text)             # 行内代码
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)       # 粗体
    text = re.sub(r"__([^_]+)__", r"\1", text)           # 粗体
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)  # 斜体
    text = text.replace("&nbsp;", " ")
    return text.strip()


def split_cell(td: str) -> list:
    """把表格单元格里的 <br> 拆成多行。"""
    return [strip_inline(x) for x in re.split(r"<br\s*/?>", td) if strip_inline(x)]


def split_row(row_line: str) -> list:
    """把 markdown 表格行拆成单元格列表；先剥外层 |，再按内部 | 拆分。"""
    s = row_line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [strip_inline(c) for c in s.split("|")]


def main() -> int:
    lines = SRC.read_text(encoding="utf-8").split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        # 代码块
        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            CODE("\n".join(buf).rstrip("\n"))
            continue

        # 表格
        if line.lstrip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|", lines[i + 1]):
            header = split_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                cells = split_row(lines[i])
                if len(cells) > len(header):
                    cells = cells[: len(header)]
                while len(cells) < len(header):
                    cells.append("")
                rows.append(cells)
                i += 1
            TABLE(header, rows)
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            H(level, strip_inline(m.group(2)))
            i += 1
            continue

        # 块引用
        b = re.match(r"^>\s?(.*)$", line)
        if b:
            # 收集连续引用行
            quote = []
            while i < n and re.match(r"^>\s?", lines[i]):
                quote.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            # 去掉引用内部的空行
            quote = [q for q in quote if q.strip()]
            # 若整段只含一个非空行，用普通段；否则列表
            body = "".join(quote)
            if len(quote) == 1 or "\n" not in body:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.line_spacing = 1.4
                r = p.add_run(strip_inline(quote[0]))
                set_run_font(r, size=10.5, color=(0x40, 0x40, 0x40))
                p.paragraph_format.space_after = Pt(6)
            else:
                for q in quote:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.8)
                    p.paragraph_format.line_spacing = 1.4
                    r = p.add_run(strip_inline(q))
                    set_run_font(r, size=10.5, color=(0x40, 0x40, 0x40))
                    p.paragraph_format.space_after = Pt(2)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
            continue

        # 列表
        ul = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        ol = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if ul or ol:
            items = []
            ordered = ol is not None
            while i < n:
                cand = re.match(r"^(\s*)[-*]\s+(.*)$", lines[i]) or re.match(r"^(\s*)\d+\.\s+(.*)$", lines[i])
                if cand is None:
                    break
                raw = cand.group(2)
                # 粗体前缀：**Label** rest...
                bm = re.match(r"^\*\*([^*]+)\*\*\s*(.*)$", raw)
                if bm:
                    items.append((strip_inline(bm.group(1)), " " + strip_inline(bm.group(2)) if bm.group(2).strip() else ""))
                else:
                    items.append(strip_inline(raw))
                i += 1
            if ordered:
                NUM(items)
            else:
                BUL(items)
            continue

        # 水平分隔线
        if re.match(r"^\s*---+\s*$", line):
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        P(strip_inline(line))
        i += 1

    doc.save(str(OUT))
    print(f"wrote {OUT}  ({len(doc.paragraphs)} paragraphs)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
